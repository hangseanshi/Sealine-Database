"""
file_generator.py — File generation tools for the Sealine Data Chat agent.

Provides three tool functions (generate_plot, generate_pdf, generate_excel)
that the Claude agent invokes to create downloadable files. Each function
generates a file, saves it to a temp directory, and returns metadata used
to emit SSE events to the React frontend.

Also exports FILE_TOOLS — the list of JSON tool definitions passed to the
Claude API — and a cleanup_expired_files utility for the background thread.

Dependencies:
    pip install matplotlib plotly openpyxl weasyprint reportlab
"""

from __future__ import annotations

import logging
import os
import re
import time
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

# Configure matplotlib for non-GUI (server) use BEFORE importing pyplot.
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
import matplotlib.colors as mcolors  # noqa: E402

import numpy as np  # noqa: E402

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Tool definitions (JSON schemas passed to Claude API)
# ---------------------------------------------------------------------------

GENERATE_PLOT_TOOL: dict[str, Any] = {
    "name": "generate_plot",
    "description": (
        "Generate a chart or plot from data. Supports bar, bar_stacked (grouped series), line, scatter, pie, "
        "heatmap, histogram, and map (interactive geographic map with OpenStreetMap "
        "tiles) chart types. ALWAYS use plot_type='map' with interactive=true when "
        "displaying geographic/location data with latitude and longitude coordinates. "
        "Use matplotlib for static charts or Leaflet.js for interactive maps / Plotly for interactive charts."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "plot_type": {
                "type": "string",
                "enum": ["bar", "bar_stacked", "line", "scatter", "pie", "heatmap", "histogram", "map"],
                "description": (
                    "Chart type. Use 'map' for any geographic/location data with "
                    "lat/lon coordinates — it renders an interactive Leaflet map. "
                    "Use 'bar_stacked' for stacked/grouped bar charts with multiple series. "
                    "Other types: bar, line, scatter, pie, heatmap, histogram."
                ),
            },
            "title": {"type": "string", "description": "Chart title."},
            "data": {
                "type": "object",
                "description": (
                    "Chart data as JSON. "
                    'For bar/line/histogram: {"labels": [...], "values": [...]}. '
                    'For bar_stacked: {"labels": [...], "series": [{"name": "SeriesA", "values": [...]}, {"name": "SeriesB", "values": [...]}]}. '
                    'For scatter: {"x": [...], "y": [...]}. '
                    'For pie: {"labels": [...], "values": [...]}. '
                    'For heatmap: {"labels_x": [...], "labels_y": [...], "values": [[...]]}. '
                    'For map with a SINGLE route: {"lat": [...], "lon": [...], "labels": [...], '
                    '"values": [...], "sizes": [...], "arrows": true/false, '
                    '"connections": [[from_idx, to_idx], ...]} '
                    'where labels are hover text (e.g. container ID), '
                    'values are optional numeric values shown on hover, '
                    '"arrows": true auto-connects points in sequence with directional arrows, '
                    '"connections" explicitly defines which point pairs to connect with arrows '
                    '(e.g. [[0,1],[1,2]] draws arrows from point 0→1 and 1→2). '
                    'For map with MULTIPLE containers/routes — ALWAYS use the GROUPS format. '
                    'Pass flat parallel arrays with a "groups" key; Python auto-groups them into separate coloured routes: '
                    '{"lat": [lat1,lat2,...], "lon": [lon1,lon2,...], '
                    '"labels": ["label1","label2",...], '
                    '"groups": ["CONTAINER_A","CONTAINER_A","CONTAINER_B","CONTAINER_B",...], '
                    '"arrows": true} '
                    'All points with the same group name form one coloured route with arrows. '
                    'Each distinct group value gets a distinct colour automatically. '
                    'DO NOT build nested route objects — use flat arrays + groups instead. '
                    '"labels" per point are popup text shown on click. '
                    'To highlight geographic zones (e.g. war zones, risk areas) as '
                    'filled polygon overlays, add a "zones" key: '
                    '{"lat": [...], "lon": [...], "labels": [...], '
                    '"zones": [{"name": "Red Sea War Zone", '
                    '"lat": [lat1, lat2, ...], "lon": [lon1, lon2, ...], '
                    '"color": "rgba(255,0,0,0.25)"}]} '
                    'Each zone is a closed polygon — repeat the first point at the end '
                    'to close the shape. Zones are rendered as semi-transparent fills. '
                    'To highlight countries or world regions by name, add '
                    '"highlight_regions": [{"name": "China", "color": "rgba(255,0,0,0.25)"}, '
                    '{"name": "United States", "color": "rgba(31,71,136,0.25)"}] '
                    'to the map data. Country names are matched case-insensitively. '
                    'Also supports ISO alpha-2 codes (e.g. "CN", "US", "DE"). '
                    'To render bubble maps where marker size reflects a numeric value, '
                    'include "sizes": [n1, n2, ...] — values are scaled to pixel radius.'
                ),
            },
            "interactive": {
                "type": "boolean",
                "description": (
                    "If true, generate Plotly HTML (required for map type). "
                    "If false, generate matplotlib PNG."
                ),
                "default": False,
            },
            "x_label": {"type": "string", "description": "X-axis label (optional, not used for map)."},
            "y_label": {"type": "string", "description": "Y-axis label (optional, not used for map)."},
        },
        "required": ["plot_type", "title", "data"],
    },
}

GENERATE_PDF_TOOL: dict[str, Any] = {
    "name": "generate_pdf",
    "description": (
        "Generate a PDF report with a title, optional summary, and data table. "
        "Use when user asks for PDF or downloadable report."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "title": {"type": "string"},
            "summary": {
                "type": "string",
                "description": "Optional summary paragraph above the table.",
            },
            "columns": {"type": "array", "items": {"type": "string"}},
            "rows": {
                "type": "array",
                "items": {"type": "array", "items": {"type": "string"}},
            },
            "filename": {
                "type": "string",
                "description": "Output filename without extension.",
            },
        },
        "required": ["title", "columns", "rows"],
    },
}

GENERATE_EXCEL_TOOL: dict[str, Any] = {
    "name": "generate_excel",
    "description": (
        "Generate a formatted Excel (.xlsx) report with blue header row "
        "(#1F4788), frozen top row, and auto-sized columns."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Sheet name."},
            "columns": {"type": "array", "items": {"type": "string"}},
            "rows": {"type": "array", "items": {"type": "array", "items": {}}},
            "filename": {
                "type": "string",
                "description": "Output filename without extension.",
            },
        },
        "required": ["title", "columns", "rows"],
    },
}

GENERATE_WORD_LABEL_TOOL: dict[str, Any] = {
    "name": "generate_word_label",
    "description": (
        "Generate a Word (.docx) shipping label document showing container events "
        "grouped by location. Each location is a bold/underlined header formatted as "
        "'LocationName/CountryCode (LOCode)'. Under each location, containers are "
        "numbered sequentially with their events listed in order. "
        "Use when the user asks for a shipping label, container label, Word label, "
        "or a formatted document showing container events by location."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "Output filename without extension.",
            },
            "locations": {
                "type": "array",
                "description": "Ordered list of locations, each with its containers and events.",
                "items": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string", "description": "Location name, e.g. 'Houston'."},
                        "country_code": {"type": "string", "description": "2-letter country code, e.g. 'US'."},
                        "locode": {"type": "string", "description": "UN/LOCODE, e.g. 'USHOU'."},
                        "containers": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "container_number": {"type": "string"},
                                    "events": {
                                        "type": "array",
                                        "description": "Events at this location, ordered by order_id ascending.",
                                        "items": {
                                            "type": "object",
                                            "properties": {
                                                "date": {"type": "string", "description": "Date in YYYY-MM-DD format."},
                                                "actual": {"type": "boolean", "description": "True if this is an actual event."},
                                                "description": {"type": "string"},
                                            },
                                            "required": ["date", "actual", "description"],
                                        },
                                    },
                                },
                                "required": ["container_number", "events"],
                            },
                        },
                    },
                    "required": ["name", "country_code", "locode", "containers"],
                },
            },
        },
        "required": ["locations"],
    },
}

FILE_TOOLS: list[dict[str, Any]] = [
    GENERATE_PLOT_TOOL,
    GENERATE_PDF_TOOL,
    GENERATE_EXCEL_TOOL,
    GENERATE_WORD_LABEL_TOOL,
]

# ---------------------------------------------------------------------------
# Shared constants
# ---------------------------------------------------------------------------

HEADER_COLOR_HEX = "#1F4788"  # Dark-blue header used for Excel & PDF
ALT_ROW_COLOR = "#F5F5F5"     # Light-gray alternating row background
MAX_COL_WIDTH = 50             # Excel column max width (chars)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _short_uuid() -> str:
    """Return the first 8 characters of a UUID4 hex string."""
    return uuid.uuid4().hex[:8]


def _slugify(text: str) -> str:
    """Convert *text* to a filesystem-safe slug (lowercase, underscores)."""
    slug = text.lower().strip()
    slug = re.sub(r"[^\w\s-]", "", slug)   # strip special chars
    slug = re.sub(r"[\s-]+", "_", slug)     # spaces / dashes -> underscores
    slug = slug.strip("_")
    return slug[:80] or "file"              # cap length, fallback


def ensure_file_store(path: str) -> None:
    """Create the file-store directory if it doesn't already exist."""
    os.makedirs(path, exist_ok=True)


def _file_meta(
    file_id: str,
    display_name: str,
    file_type: str,
    file_path: str,
) -> dict[str, Any]:
    """Build the standard metadata dict returned by every generator."""
    size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
    return {
        "file_id": file_id,
        "filename": display_name,
        "file_type": file_type,
        "file_path": file_path,
        "size_bytes": size,
    }


def _is_numeric(value: Any) -> bool:
    """Return True if *value* looks like a number (int, float, or numeric string)."""
    if isinstance(value, (int, float)):
        return True
    if isinstance(value, str):
        try:
            float(value.replace(",", ""))
            return True
        except (ValueError, AttributeError):
            return False
    return False


# ---------------------------------------------------------------------------
# 1. generate_plot
# ---------------------------------------------------------------------------


def generate_plot(
    plot_type: str,
    title: str,
    data: dict[str, Any],
    interactive: bool = False,
    x_label: str | None = None,
    y_label: str | None = None,
    file_store_path: str = "./tmp/files",
) -> dict[str, Any]:
    """Generate a chart and save it to *file_store_path*.

    Returns a metadata dict consumed by the SSE layer.
    """
    ensure_file_store(file_store_path)
    file_id = _short_uuid()
    title_slug = _slugify(title)

    try:
        if interactive:
            return _plot_interactive(
                plot_type, title, data, x_label, y_label,
                file_id, title_slug, file_store_path,
            )
        return _plot_static(
            plot_type, title, data, x_label, y_label,
            file_id, title_slug, file_store_path,
        )
    except Exception as exc:
        logger.exception("generate_plot failed")
        return {"error": str(exc)}


# -- static (matplotlib) ---------------------------------------------------


def _plot_static(
    plot_type: str,
    title: str,
    data: dict[str, Any],
    x_label: str | None,
    y_label: str | None,
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Render a chart with matplotlib and save as PNG."""

    # Use a clean built-in style.  seaborn-v0_8-whitegrid ships with
    # matplotlib >= 3.6.  Fall back gracefully for older versions.
    try:
        plt.style.use("seaborn-v0_8-whitegrid")
    except OSError:
        try:
            plt.style.use("seaborn-whitegrid")
        except OSError:
            pass  # use default style

    fig, ax = plt.subplots(figsize=(10, 6))

    labels = data.get("labels", [])
    values = data.get("values", [])

    if plot_type == "bar":
        x_positions = range(len(labels))
        ax.bar(x_positions, values, color=HEADER_COLOR_HEX, edgecolor="white")
        ax.set_xticks(list(x_positions))
        ax.set_xticklabels(labels, rotation=45, ha="right")

    elif plot_type == "line":
        ax.plot(labels, values, marker="o", linewidth=2, color=HEADER_COLOR_HEX)
        ax.tick_params(axis="x", rotation=45)

    elif plot_type == "scatter":
        x = data.get("x", [])
        y = data.get("y", [])
        ax.scatter(x, y, alpha=0.7, color=HEADER_COLOR_HEX, edgecolors="white")

    elif plot_type == "pie":
        # Align labels/values to the shorter array, then cap at top 20 slices
        n = min(len(labels), len(values))
        labels, values = list(labels[:n]), list(values[:n])
        PIE_MAX = 20
        if n > PIE_MAX:
            # Sort by value descending, keep top PIE_MAX-1, group rest as Other
            paired = sorted(zip(values, labels), reverse=True)
            top_vals  = [v for v, _ in paired[:PIE_MAX - 1]]
            top_lbls  = [l for _, l in paired[:PIE_MAX - 1]]
            other_val = sum(v for v, _ in paired[PIE_MAX - 1:])
            values = top_vals + [other_val]
            labels = top_lbls + ["Other"]
        # Pie uses fig-level; close the original figure to prevent leak,
        # then create a fresh one.
        plt.close(fig)
        fig, ax = plt.subplots(figsize=(10, 6))
        wedges, texts, autotexts = ax.pie(
            values,
            labels=labels,
            autopct="%1.1f%%",
            startangle=140,
            textprops={"fontsize": 10},
        )
        ax.set_aspect("equal")

    elif plot_type == "heatmap":
        labels_x = data.get("labels_x", [])
        labels_y = data.get("labels_y", [])
        matrix = np.array(values, dtype=float)
        im = ax.imshow(matrix, cmap="Blues", aspect="auto")
        ax.set_xticks(range(len(labels_x)))
        ax.set_xticklabels(labels_x, rotation=45, ha="right")
        ax.set_yticks(range(len(labels_y)))
        ax.set_yticklabels(labels_y)
        fig.colorbar(im, ax=ax)

    elif plot_type == "histogram":
        ax.hist(values, bins="auto", color=HEADER_COLOR_HEX, edgecolor="white")

    elif plot_type == "bar_stacked":
        series_list = data.get("series", [])
        x_pos = range(len(labels))
        bottom = [0.0] * len(labels)
        PALETTE = [
            "#1F4788", "#E07B39", "#2ECC71", "#E74C3C", "#9B59B6",
            "#F39C12", "#1ABC9C", "#E91E63", "#00BCD4", "#8BC34A",
        ]
        for idx, s in enumerate(series_list):
            s_vals = s.get("values", [0] * len(labels))
            ax.bar(
                x_pos, s_vals, bottom=bottom,
                label=s.get("name", f"Series {idx+1}"),
                color=PALETTE[idx % len(PALETTE)],
                edgecolor="white",
            )
            bottom = [b + v for b, v in zip(bottom, s_vals)]
        ax.set_xticks(list(x_pos))
        ax.set_xticklabels(labels, rotation=45, ha="right")
        ax.legend(fontsize=9)

    else:
        plt.close(fig)
        return {"error": f"Unsupported plot_type: {plot_type}"}

    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    if x_label:
        ax.set_xlabel(x_label)
    if y_label:
        ax.set_ylabel(y_label)

    fig.tight_layout()

    filename = f"{file_id}_{title_slug}.png"
    full_path = os.path.join(file_store_path, filename)
    fig.savefig(full_path, dpi=150, bbox_inches="tight")
    plt.close(fig)

    return _file_meta(file_id, f"{title_slug}.png", "image/png", full_path)


# ---------------------------------------------------------------------------
# Leaflet map HTML template (placeholders: |||TITLE|||, |||DATA_JSON|||,
# |||CENTER_LAT|||, |||CENTER_LON|||, |||ZOOM|||)
# ---------------------------------------------------------------------------
_LEAFLET_MAP_HTML = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>|||TITLE|||</title>
<script>
window.onerror=function(m,s,l,c,e){
  var el=document.getElementById('map-error');
  if(el){el.style.display='block';el.innerHTML='<b>JS Error:</b> '+m+' (line '+l+')';}
  else if(document.body){document.body.style.background='#fff';document.body.innerHTML='<pre style="color:red;padding:20px">JS Error: '+m+'\\nLine: '+l+'</pre>';}
  else{setTimeout(function(){document.body&&(document.body.style.background='#fff',document.body.innerHTML='<pre style="color:red;padding:20px">JS Error: '+m+'\\nLine: '+l+'</pre>');},200);}
  return false;
};
</script>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script src="https://cdn.jsdelivr.net/npm/topojson-client@3/dist/topojson-client.min.js"></script>
<style>
  html, body, #map { margin: 0; padding: 0; width: 100%; height: 100vh; overflow: hidden; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }
  .map-title {
    position: absolute; top: 12px; left: 50%; transform: translateX(-50%);
    z-index: 1000; background: rgba(255,255,255,0.95);
    padding: 7px 20px; border-radius: 8px;
    font-size: 15px; font-weight: 700; color: #1F4788;
    box-shadow: 0 2px 10px rgba(0,0,0,0.18);
    white-space: nowrap; pointer-events: none;
  }
  .legend {
    background: rgba(255,255,255,0.95); border-radius: 8px;
    padding: 10px 14px; box-shadow: 0 2px 10px rgba(0,0,0,0.18);
    max-height: 280px; overflow-y: auto; min-width: 150px;
  }
  .legend h4 {
    margin: 0 0 8px 0; font-size: 11px; font-weight: 700;
    color: #555; text-transform: uppercase; letter-spacing: 0.6px;
  }
  .legend-item { display: flex; align-items: center; gap: 8px; margin-bottom: 6px; font-size: 12px; color: #333; }
  .leg-line { width: 22px; height: 3px; border-radius: 2px; flex-shrink: 0; }
  .leg-dot { width: 11px; height: 11px; border-radius: 50%; border: 2px solid rgba(255,255,255,0.8); flex-shrink: 0; }
  .leg-zone { width: 14px; height: 10px; border-radius: 3px; flex-shrink: 0; }
  .route-tooltip { background: rgba(255,255,255,0.97); border: 1px solid #ccc; border-radius: 5px; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }
  .stop-label {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    font-size: 11px; font-weight: 600;
    color: #1a1a1a;
    white-space: nowrap;
    text-shadow: 1px 1px 0 #fff, -1px -1px 0 #fff, 1px -1px 0 #fff, -1px 1px 0 #fff;
    pointer-events: none;
  }
  .leaflet-popup-content { font-size: 11px; line-height: 1.5; min-width: 320px; white-space: nowrap; }
</style>
</head>
<body>
<div class="map-title">|||TITLE|||</div>
<div id="map"></div>
<div id="map-error" style="display:none;position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);
  background:#fff;padding:20px;border:2px solid red;border-radius:8px;z-index:9999;
  font-family:monospace;font-size:13px;max-width:80%;word-break:break-all;"></div>
<script>
try {
// ── Data ─────────────────────────────────────────────────────────────────
const MAP_DATA = |||DATA_JSON|||;

// ── Country name → ISO numeric code mapping ───────────────────────────────
const COUNTRY_CODES = {
  'afghanistan':4,'albania':8,'algeria':12,'angola':24,'argentina':32,
  'australia':36,'austria':40,'bahrain':48,'bangladesh':50,'belgium':56,
  'bolivia':68,'brazil':76,'bulgaria':100,'cambodia':116,'cameroon':120,
  'canada':124,'chile':152,'china':156,'prc':156,'cn':156,'colombia':170,
  'congo':180,'costa rica':188,'croatia':191,'cuba':192,'czech republic':203,
  'czechia':203,'denmark':208,'dk':208,'dominican republic':214,
  'ecuador':218,'egypt':818,'eg':818,'ethiopia':231,'finland':246,'fi':246,
  'france':250,'fr':250,'germany':276,'de':276,'ghana':288,'greece':300,
  'gr':300,'hong kong':344,'hk':344,'hungary':348,'india':356,'in':356,
  'indonesia':360,'id':360,'iran':364,'iraq':368,'ireland':372,'israel':376,
  'il':376,'italy':380,'it':380,'jamaica':388,'japan':392,'jp':392,
  'jordan':400,'kenya':404,'south korea':410,'korea':410,'kr':410,
  'kuwait':414,'kw':414,'laos':418,'latvia':428,'libya':434,'malaysia':458,
  'my':458,'mexico':484,'mx':484,'morocco':504,'ma':504,'mozambique':508,
  'myanmar':104,'burma':104,'mm':104,'netherlands':528,'nl':528,'holland':528,
  'new zealand':554,'nz':554,'nicaragua':558,'nigeria':566,'ng':566,
  'norway':578,'no':578,'oman':512,'om':512,'pakistan':586,'pk':586,
  'panama':591,'pa':591,'peru':604,'philippines':608,'ph':608,'poland':616,
  'pl':616,'portugal':620,'pt':620,'qatar':634,'qa':634,'romania':642,
  'russia':643,'ru':643,'saudi arabia':682,'sa':682,'senegal':686,
  'singapore':702,'sg':702,'somalia':706,'south africa':710,'za':710,
  'spain':724,'es':724,'sri lanka':144,'lk':144,'sudan':729,'sd':729,
  'sweden':752,'se':752,'switzerland':756,'ch':756,'taiwan':158,'tw':158,
  'tanzania':834,'tz':834,'thailand':764,'th':764,'tunisia':788,'tn':788,
  'turkey':792,'tr':792,'ukraine':804,'ua':804,'united arab emirates':784,
  'uae':784,'ae':784,'united kingdom':826,'uk':826,'gb':826,'britain':826,
  'united states':840,'usa':840,'us':840,'america':840,'uruguay':858,
  'venezuela':862,'vietnam':704,'vn':704,'yemen':887,'ye':887,
  'zambia':894,'zimbabwe':716
};

// ── Init map — CartoDB Voyager (English labels everywhere) ────────────────
if (!window.L) throw new Error('Leaflet not loaded — window.L is ' + typeof window.L);
const map = L.map('map', { zoomControl: true, scrollWheelZoom: true })
             .setView([|||CENTER_LAT|||, |||CENTER_LON|||], |||ZOOM|||);
// Force re-layout after iframe finishes sizing (fixes blank-map-in-iframe bug)
// Call multiple times to cover browsers that finish sizing at different times
[100, 300, 600, 1200].forEach(function(ms) {
  setTimeout(function() { map.invalidateSize(false); }, ms);
});
window.addEventListener('resize', function() { map.invalidateSize(false); });

L.tileLayer('https://{s}.basemaps.cartocdn.com/rastertiles/voyager/{z}/{x}/{y}{r}.png', {
  attribution: '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> &copy; <a href="https://carto.com/">CARTO</a>',
  subdomains: 'abcd', maxZoom: 19
}).addTo(map);

// ── Helpers ───────────────────────────────────────────────────────────────
function parseRgba(str) {
  if (!str) return { r: 31, g: 71, b: 136, a: 0.25 };
  const m = str.match(/rgba?\\((\\d+),\\s*(\\d+),\\s*(\\d+)(?:,\\s*([\\d.]+))?\\)/);
  if (m) return { r: +m[1], g: +m[2], b: +m[3], a: m[4] !== undefined ? +m[4] : 1 };
  if (str[0] === '#') {
    const h = str.slice(1);
    return { r: parseInt(h.slice(0,2),16), g: parseInt(h.slice(2,4),16), b: parseInt(h.slice(4,6),16), a: 1 };
  }
  return { r: 31, g: 71, b: 136, a: 0.25 };
}
function toHex(r,g,b) {
  return '#' + [r,g,b].map(v => v.toString(16).padStart(2,'0')).join('');
}

// ── Label helpers ─────────────────────────────────────────────────────────
function extractLabelInfo(lbl) {
  // Location: first <br>-delimited segment, strip HTML tags
  var segs = lbl.split(/<br\\s*\\/?>/i);
  var loc = (segs[0] || '').replace(/<[^>]+>/g, '').trim();
  // Date: first YYYY-MM-DD (A/E) pattern in the label
  var dm = lbl.match(/(\\d{4}-\\d{2}-\\d{2}\\s*\\([AE]\\))/);
  return { loc: loc, date: dm ? dm[1] : '' };
}

function extractLabelInfoFor(lbl, containerName) {
  var segs = lbl.split(/<br\\s*\\/?>/i);
  var loc = (segs[0] || '').replace(/<[^>]+>/g, '').trim();
  // Extract tracking number from data-trk attribute embedded in the label
  var trkM = lbl.match(/data-trk="([^"]+)"/);
  var tracking = trkM ? trkM[1] : '';
  // If multiple containers share this stop, find the date for the specific container
  if (containerName) {
    var esc = containerName.replace(/[.*+?^${}()|[\\]\\\\]/g, '\\\\$&');
    var cm = lbl.match(new RegExp(esc + '[^\\\\d]*(\\\\d{4}-\\\\d{2}-\\\\d{2}\\\\s*\\\\([AE]\\\\))'));
    if (cm) return { loc: loc, date: cm[1], tracking: tracking };
  }
  var dm = lbl.match(/(\\d{4}-\\d{2}-\\d{2}\\s*\\([AE]\\))/);
  return { loc: loc, date: dm ? dm[1] : '', tracking: tracking };
}

// ── Pure-Leaflet arrow line (no plugin needed) ────────────────────────────
function drawArrowLine(fromPt, toPt, color, laneOffset, tooltipHtml) {
  // Pure Mercator math — no dependency on map state (avoids NaN from map.project)
  laneOffset = laneOffset || 0;
  var WSIZ = 1024; // world pixels at zoom 2 (256 * 2^2)
  function merc(lat, lon) {
    var s = Math.sin(lat * Math.PI / 180);
    return {
      x: (lon + 180) / 360 * WSIZ,
      y: (0.5 - Math.log((1 + s) / (1 - s)) / (4 * Math.PI)) * WSIZ
    };
  }
  function unmerc(px, py) {
    var n = Math.PI - 2 * Math.PI * py / WSIZ;
    return { lat: 180 / Math.PI * Math.atan(0.5 * (Math.exp(n) - Math.exp(-n))),
             lng: px / WSIZ * 360 - 180 };
  }
  var p1 = merc(fromPt[0], fromPt[1]);
  var p2 = merc(toPt[0],   toPt[1]);
  var dx = p2.x - p1.x, dy = p2.y - p1.y;
  var pixChord = Math.sqrt(dx * dx + dy * dy) || 1;

  // When stops are very close together the laneOffset (in absolute pixel space)
  // would dwarf the chord, swinging the bezier control point wildly and
  // producing huge crossing arcs.  Use a plain straight line for short segments.
  var STRAIGHT_THRESHOLD = 10; // pixels at WSIZ=1024 ≈ 400 km real distance
  if (pixChord < STRAIGHT_THRESHOLD) {
    var _sl = L.polyline([fromPt, toPt], { color: color, weight: 3, opacity: 0.85 }).addTo(map);
    if (tooltipHtml) _sl.bindTooltip(tooltipHtml, { sticky: true, className: 'route-tooltip' });
    var mLat = (fromPt[0] + toPt[0]) / 2;
    var mLon = (fromPt[1] + toPt[1]) / 2;
    var mAngle = Math.atan2(toPt[1] - fromPt[1], toPt[0] - fromPt[0]) * 180 / Math.PI;
    var mSvg = '<svg width="18" height="18" viewBox="-9 -9 18 18" xmlns="http://www.w3.org/2000/svg">'
      + '<polygon points="0,-7 5,3 0,0 -5,3" fill="' + color + '" opacity="0.95"'
      + ' transform="rotate(' + mAngle.toFixed(1) + ')"/></svg>';
    L.marker([mLat, mLon], {
      icon: L.divIcon({ html: mSvg, className: '', iconSize: [18, 18], iconAnchor: [9, 9] }),
      interactive: false, zIndexOffset: 100,
    }).addTo(map);
    return;
  }

  // Perpendicular unit vector (left of route direction) for lane separation
  var perpX = -dy / pixChord;
  var perpY =  dx / pixChord;

  // Control point: northward by 20% of chord (fully proportional — no fixed minimum
  // so nearby stops don't produce huge spurious arcs).
  var pxOffset = pixChord * 0.20;
  var ctrl     = unmerc(p1.x + dx / 2 + perpX * laneOffset,
                        p1.y + dy / 2 - pxOffset + perpY * laneOffset);
  var ctrlLat  = Math.min(85, Math.max(-85, ctrl.lat));
  var ctrlLon  = ctrl.lng;

  // Sample 50 points along the quadratic bezier in lat/lng space
  const N = 50;
  const curvePts = [];
  for (let i = 0; i <= N; i++) {
    const t = i / N, u = 1 - t;
    curvePts.push([
      u*u*fromPt[0] + 2*u*t*ctrlLat + t*t*toPt[0],
      u*u*fromPt[1] + 2*u*t*ctrlLon + t*t*toPt[1],
    ]);
  }
  var _cl = L.polyline(curvePts, { color: color, weight: 3, opacity: 0.85 }).addTo(map);
  if (tooltipHtml) _cl.bindTooltip(tooltipHtml, { sticky: true, className: 'route-tooltip' });

  // Arrow at t=0.75 (descending leg toward destination, away from arc apex)
  function bzPt(t) {
    const u = 1 - t;
    return [u*u*fromPt[0]+2*u*t*ctrlLat+t*t*toPt[0], u*u*fromPt[1]+2*u*t*ctrlLon+t*t*toPt[1]];
  }
  const arrPt = bzPt(0.75);
  const arrowLat = arrPt[0], arrowLon = arrPt[1];
  const pa = bzPt(0.73), pb = bzPt(0.77);
  const angle = Math.atan2(pb[1] - pa[1], pb[0] - pa[0]) * 180 / Math.PI;

  const svg = '<svg width="18" height="18" viewBox="-9 -9 18 18" xmlns="http://www.w3.org/2000/svg">'
    + '<polygon points="0,-7 5,3 0,0 -5,3" fill="' + color + '" opacity="0.95"'
    + ' transform="rotate(' + angle.toFixed(1) + ')"/></svg>';
  L.marker([arrowLat, arrowLon], {
    icon: L.divIcon({ html: svg, className: '', iconSize: [18, 18], iconAnchor: [9, 9] }),
    interactive: false, zIndexOffset: 100,
  }).addTo(map);
}

// ── Legend ────────────────────────────────────────────────────────────────
const legendCtrl = L.control({ position: 'bottomright' });
legendCtrl.onAdd = function() {
  const div = L.DomUtil.create('div', 'legend');
  let html = '<h4>Legend</h4>';
  MAP_DATA.routes.forEach(r => {
    if (!r.name) return;
    const hasLines = r.connections && r.connections.length > 0;
    html += '<div class="legend-item">';
    if (hasLines) html += '<div class="leg-line" style="background:' + r.color + '"></div>';
    html += '<div class="leg-dot" style="background:' + r.color + '"></div>';
    html += '<span>' + r.name + '</span></div>';
  });
  MAP_DATA.zones.forEach(z => {
    const c = parseRgba(z.color);
    const hex = toHex(c.r, c.g, c.b);
    html += '<div class="legend-item"><div class="leg-zone" style="background:' + hex
          + ';opacity:0.55;border:2px solid ' + hex + '"></div><span>' + z.name + '</span></div>';
  });
  (MAP_DATA.highlight_regions || []).forEach(r => {
    const c = parseRgba(r.color);
    const hex = toHex(c.r, c.g, c.b);
    html += '<div class="legend-item"><div class="leg-zone" style="background:' + hex
          + ';opacity:0.55;border:2px solid ' + hex + '"></div><span>' + r.name + '</span></div>';
  });
  div.innerHTML = html;
  return div;
};
legendCtrl.addTo(map);

// ── Zone polygons ──────────────────────────────────────────────────────────
MAP_DATA.zones.forEach(function(zone) {
  const c = parseRgba(zone.color);
  const hex = toHex(c.r, c.g, c.b);
  L.polygon(zone.latlngs, {
    color: hex, weight: 2, opacity: 0.85,
    fillColor: hex, fillOpacity: c.a,
  }).addTo(map).bindTooltip('<b>' + zone.name + '</b>', { sticky: true });
});

// ── Country / region highlights (TopoJSON world atlas) ────────────────────
const highlightRegions = MAP_DATA.highlight_regions || [];
if (highlightRegions.length > 0) {
  const codeColorMap = {};
  highlightRegions.forEach(function(region) {
    const key = (region.name || '').toLowerCase().trim();
    const numCode = COUNTRY_CODES[key];
    if (numCode !== undefined) {
      codeColorMap[numCode] = { color: region.color, name: region.name };
    }
  });
  fetch('https://cdn.jsdelivr.net/npm/world-atlas@2/countries-110m.json')
    .then(function(r) { return r.json(); })
    .then(function(world) {
      if (typeof topojson === 'undefined') return;
      const countries = topojson.feature(world, world.objects.countries);
      L.geoJSON(countries, {
        style: function(feature) {
          const match = codeColorMap[+feature.id];
          if (match) {
            const c = parseRgba(match.color);
            const hex = toHex(c.r, c.g, c.b);
            return { color: hex, weight: 1.5, opacity: 0.8, fillColor: hex, fillOpacity: c.a };
          }
          return { fillOpacity: 0, opacity: 0, weight: 0 };
        },
        onEachFeature: function(feature, layer) {
          const match = codeColorMap[+feature.id];
          if (match) layer.bindTooltip('<b>' + match.name + '</b>', { sticky: true });
        }
      }).addTo(map);
    }).catch(function() {});
}

// ── Routes — lines + arrows + labelled dots ───────────────────────────────
MAP_DATA.routes.forEach(function(route) {
  const pts = route.points;        // [[lat, lon], ...]
  const color = route.color;
  const connections = route.connections || [];
  const sizes = route.sizes || [];
  const labels = route.labels || [];

  if (!pts || pts.length === 0) return;

  // 1. Draw lines with arrow markers + hover tooltips
  const laneOffsets = route.lane_offsets || [];
  const rName = route.name || '';
  const rTracking = route.tracking || '';
  connections.forEach(function(conn, ci) {
    const fi = conn[0], ti = conn[1];
    if (fi >= pts.length || ti >= pts.length) return;
    var tip;
    if (!MAP_DATA.hide_route_tracking) {
      // Tracking route map — label segments: [TrackNumber, Location/Country (LOCode), RouteType, Date, ...]
      var fromSegs = (labels[fi] || '').split(/<br\s*\/?>/i);
      var toSegs   = (labels[ti] || '').split(/<br\s*\/?>/i);
      var fromLoc  = (fromSegs[1] || '?').replace(/<[^>]+>/g, '').trim();
      var toLoc    = (toSegs[1]   || '?').replace(/<[^>]+>/g, '').trim();
      var fromType = (fromSegs[2] || '').replace(/<[^>]+>/g, '').trim().toUpperCase();
      var toType   = (toSegs[2]   || '').replace(/<[^>]+>/g, '').trim().toUpperCase();
      var fromDate = (fromSegs[3] || '').replace(/<[^>]+>/g, '').trim();
      var toDate   = (toSegs[3]   || '').replace(/<[^>]+>/g, '').trim();
      // Strip the trailing LOCode "(XXYYY)" from location names so tooltip shows
      // "City/Country (ROUTETYPE)" instead of "City/Country (LOCODE) (ROUTETYPE)".
      var fromLocClean = fromLoc.replace(/\\s*\\([A-Z0-9]{2,8}\\)\\s*$/, '').trim() || fromLoc;
      var toLocClean   = toLoc.replace(/\\s*\\([A-Z0-9]{2,8}\\)\\s*$/, '').trim() || toLoc;
      tip = '<div style="font-size:11px;line-height:1.6;white-space:nowrap;padding:2px 4px">'
        + '<b>' + rName + '</b><br>'
        + fromLocClean + (fromType ? ' (' + fromType + ')' : '') + ' &rarr; ' + toLocClean + (toType ? ' (' + toType + ')' : '')
        + '<br>' + (fromDate || '?') + ' &rarr; ' + (toDate || '?');
      tip += '</div>';
    } else {
      // Container route map — label segments: [ContainerNumber, Location, EventLines...]
      var fromInfo = extractLabelInfoFor(labels[fi] || '', rName);
      var toInfo   = extractLabelInfoFor(labels[ti] || '', rName);
      tip = '<div style="font-size:11px;line-height:1.6;white-space:nowrap;padding:2px 4px">'
        + '<b>' + rName + '</b><br>'
        + (fromInfo.loc || '?') + ' &rarr; ' + (toInfo.loc || '?');
      if (fromInfo.date || toInfo.date) {
        tip += '<br>' + (fromInfo.date || '?') + ' &rarr; ' + (toInfo.date || '?')
             + '&nbsp;&nbsp;<span style="color:#666;font-size:10px">(A)=Actual&nbsp;(E)=Estimated</span>';
      }
      tip += '</div>';
    }
    drawArrowLine(pts[fi], pts[ti], color, laneOffsets[ci] || 0, tip);
  });

  // 2. Draw stop markers + permanent labels
  pts.forEach(function(pt, i) {
    const rawLabel = (labels[i] || '').toString();
    // Build popup table: location header + per-container rows with event lines.
    const popupHtml = (function(raw) {
      var blocks = raw.split(/<br\\s*\\/?>/i);
      if (blocks.length < 2) return raw.replace(/\\n|\\\\n/g, '<br>');
      var header = blocks[0].replace(/<[^>]*>/g, '').trim();
      var seqRe = /^(?:(\\d+)\\.\\s+)?(\\S+)\\s*([\\s\\S]*)/;
      var rows = '';
      if (header) {
        rows += '<tr><td colspan="2" style="font-weight:bold;text-decoration:underline;text-align:center;padding-bottom:4px;white-space:nowrap">' + header + '</td></tr>';
      }
      for (var b = 1; b < blocks.length; b++) {
        var blockText = blocks[b].replace(/<[^>]*>/g, '').trim();
        if (!blockText) continue;
        var lines = blockText.split(/\\n|\\\\n/);
        var m = seqRe.exec(lines[0]);
        if (!m) continue;
        var seq = m[1] || '', container = m[2], rest = m[3].trim();
        rows += '<tr>'
          + (seq ? '<td style="padding-right:6px;font-weight:bold;vertical-align:top;white-space:nowrap">' + seq + ':</td>' : '<td></td>')
          + '<td style="font-weight:bold;white-space:nowrap">' + container + '</td>'
          + '</tr>';
        var events = [];
        if (rest) events.push(rest);
        for (var j = 1; j < lines.length; j++) {
          var ev = lines[j].trim();
          if (ev) events.push(ev);
        }
        for (var k = 0; k < events.length; k++) {
          rows += '<tr><td></td>'
            + '<td style="white-space:nowrap">&nbsp;&nbsp;|&mdash;&mdash;&mdash;&mdash;&mdash;&mdash; ' + events[k] + '</td>'
            + '</tr>';
        }
      }
      var rowCount = (rows.match(/<\\/tr>/g) || []).length;
      var table = '<table style="border-collapse:collapse;font-size:11px;line-height:1.2">' + rows + '</table>';
      if (rowCount > 10) {
        return '<div style="max-height:170px;overflow-y:auto;overflow-x:auto">' + table + '</div>';
      }
      return table;
    })(rawLabel);
    // First line only for the always-visible dot label — strip tags
    const firstLine = rawLabel.split('<br>')[0].replace(/<[^>]+>/g, '').trim();
    const radius = (sizes[i] && +sizes[i] > 0) ? Math.max(6, Math.min(30, +sizes[i])) : 10;

    // Filled circle marker
    const marker = L.circleMarker(pt, {
      radius: radius,
      color: '#ffffff', weight: 2,
      fillColor: color, fillOpacity: 1,
    }).addTo(map);

    // Popup on click — shows full detail (type, location, date, actual/estimated)
    if (popupHtml) {
      marker.bindPopup('<div style="font-size:11px;line-height:1.2;white-space:nowrap">' + popupHtml + '</div>', {maxWidth: 700});
    }

    // Permanent label above the dot — first line only (keeps the map uncluttered)
    if (firstLine) {
      const labelIcon = L.divIcon({
        html: '<div class="stop-label">' + firstLine + '</div>',
        className: '',
        iconAnchor: [0, radius + 2],
      });
      L.marker(pt, { icon: labelIcon, interactive: false, zIndexOffset: 200 }).addTo(map);
    }
  });
});

// ── Fit bounds ─────────────────────────────────────────────────────────────
const allPts = [
  ...(MAP_DATA.routes || []).flatMap(function(r) { return r.points || []; }),
  ...(MAP_DATA.zones  || []).flatMap(function(z) { return z.latlngs || []; }),
];
if (allPts.length > 0) {
  try { map.fitBounds(L.latLngBounds(allPts), { padding: [60, 60], maxZoom: 12 }); }
  catch(e) {}
}
} catch(e) {
  var d = document.getElementById('map-error');
  if (d) { d.style.display = 'block'; d.innerHTML = '<b>Map Error:</b> ' + e.message + '<br><pre>' + e.stack + '</pre>'; }
}
</script>
</body>
</html>
"""


def _plot_tracking_route_map(
    title: str,
    data: dict[str, Any],
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Generate a Leaflet.js tracking route map as a self-contained HTML file.

    Uses v_sealine_tracking_route data with structured locations and routes.
    """
    import json as _json
    import html as _html

    _TRACKING_ROUTE_HTML = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>|||TITLE|||</title>
<script>
window.onerror=function(m,s,l,c,e){
  var el=document.getElementById('map-error');
  if(el){el.style.display='block';el.innerHTML='<b>JS Error:</b> '+m+' (line '+l+')';}
  return false;
};
</script>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script src="https://cdn.jsdelivr.net/npm/topojson-client@3/dist/topojson-client.min.js"></script>
<style>
  html, body, #map { margin: 0; padding: 0; width: 100%; height: 100vh; overflow: hidden; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }
  .map-title {
    position: absolute; top: 12px; left: 50%; transform: translateX(-50%);
    z-index: 1000; background: rgba(255,255,255,0.95);
    padding: 7px 20px; border-radius: 8px;
    font-size: 15px; font-weight: 700; color: #1F4788;
    box-shadow: 0 2px 10px rgba(0,0,0,0.18);
    white-space: nowrap; pointer-events: none;
  }
  .legend {
    background: rgba(255,255,255,0.97); padding: 10px 14px;
    border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.18);
    font-size: 12px; line-height: 1.7; max-height: 300px; overflow-y: auto; min-width: 120px;
  }
  .legend h4 { margin: 0 0 6px 0; font-size: 13px; color: #1F4788; }
  .legend-item { display: flex; align-items: center; gap: 7px; padding: 1px 4px; border-radius: 3px; cursor: pointer; }
  .legend-item:hover { background: #f0f4ff; }
  .leg-line { width: 22px; height: 3px; border-radius: 2px; flex-shrink: 0; }
  .leg-dot { width: 11px; height: 11px; border-radius: 50%; border: 2px solid rgba(255,255,255,0.8); flex-shrink: 0; }
  .leg-zone { width: 22px; height: 11px; border-radius: 3px; flex-shrink: 0; }
  .route-tooltip { background: rgba(255,255,255,0.97); border: 1px solid #ccc; border-radius: 5px; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }
  .stop-label {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    font-size: 10px;
    font-weight: 600;
    color: #1F4788;
    white-space: nowrap;
  }
  #map-error {
    display: none; position: absolute; top: 60px; left: 50%; transform: translateX(-50%);
    z-index: 9999; background: #fff3cd; border: 1px solid #ffc107;
    padding: 10px 16px; border-radius: 6px; font-size: 13px;
  }
</style>
</head>
<body>
<div id="map"></div>
<div class="map-title">|||TITLE|||</div>
<div id="map-error"></div>
<script>
try {
var ROUTE_DATA = |||ROUTE_DATA_JSON|||;

var map = L.map('map', {preferCanvas: true}).setView([20, 0], 2);
// CartoDB Voyager — renders all labels in English
L.tileLayer('https://{s}.basemaps.cartocdn.com/rastertiles/voyager/{z}/{x}/{y}{r}.png', {
  attribution: '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors &copy; <a href="https://carto.com/attributions">CARTO</a>',
  subdomains: 'abcd',
  maxZoom: 20
}).addTo(map);

// ── War Zone overlays ──────────────────────────────────────────────────────
var WAR_ZONES = [
  {
    // JWC JWLA-033 / UKMTO: full southern Red Sea (south of 18°N),
    // Bab el-Mandeb strait, entire Gulf of Aden, extending east to ~60°E
    name: "Red Sea / Gulf of Aden / Bab el-Mandeb (Houthi High Risk Zone)",
    coords: [
      // Northern Red Sea crossing at 18°N (Eritrea/Sudan → Saudi coast)
      [18.0,37.8],[18.0,41.0],
      // Saudi/Yemen coast going south to Bab el-Mandeb
      [16.0,43.2],[13.5,43.5],[12.5,43.5],
      // Gulf of Aden north coast (Yemen) going east
      [12.0,46.0],[12.5,49.0],[13.5,51.5],[16.5,53.0],
      // Extended east per JWC JWLA-033 (~60°E)
      [13.5,57.0],[11.0,60.0],
      // South open ocean / Somalia
      [6.0,50.0],[-1.5,41.5],
      // Back up African (Somalia/Djibouti/Eritrea) coast
      [3.0,42.0],[7.0,41.5],[10.5,42.5],
      [11.5,43.2],[14.0,41.0],[16.5,39.5],[18.0,37.8]
    ]
  },
  {
    // Gaza Strip maritime zone + Israeli coastal waters (Eastern Mediterranean)
    name: "Gaza / Israel (Conflict Zone)",
    coords: [
      [29.2,33.8],[29.2,36.5],
      [33.5,36.5],[33.5,33.8]
    ]
  },
  {
    // JWC JWLA-033: full Black Sea — mine threat, USV/drone attacks,
    // Ukrainian/Russian mutual strikes, includes Sea of Azov
    name: "Ukraine / Black Sea (Conflict Zone)",
    coords: [
      [41.0,28.0],[43.5,28.0],
      [46.5,30.0],[47.0,32.0],
      [47.1,35.0],[46.5,38.0],
      [47.5,38.5],[47.0,39.5],
      [45.5,41.5],[43.5,41.5],
      [41.5,41.5],[41.0,40.0],
      [41.0,28.0]
    ]
  },
  {
    // Sudan civil war (SAF vs RSF, April 2023–present): Khartoum, Darfur,
    // and Red Sea coast (Port Sudan)
    name: "Sudan (Civil War Zone)",
    coords: [
      [12.0,22.5],[22.0,22.5],
      [22.0,37.5],[18.5,40.0],
      [15.5,38.5],[12.5,36.5],
      [10.0,33.0],[10.0,23.5],
      [12.0,22.5]
    ]
  }
];
var warZoneLayers = WAR_ZONES.map(function(z) {
  var poly = L.polygon(z.coords, {
    color: '#c0392b', weight: 1.5, opacity: 0.85,
    fillColor: '#e74c3c', fillOpacity: 0.18, dashArray: '6,4',
    interactive: true
  }).addTo(map);
  poly.bindTooltip(
    '<div style="font-weight:bold;color:#c0392b;font-size:12px;white-space:nowrap;">&#9888; ' + z.name + '</div>',
    {sticky: true, className: 'route-tooltip'}
  );
  return poly;
});
var warZonesVisible = true;

// ── Country / region highlights (TopoJSON world atlas) ─────────────────────
var COUNTRY_CODES_TRK = {
  'afghanistan':4,'albania':8,'algeria':12,'angola':24,'argentina':32,
  'australia':36,'austria':40,'bahrain':48,'bangladesh':50,'belgium':56,
  'bolivia':68,'brazil':76,'bulgaria':100,'cambodia':116,'cameroon':120,
  'canada':124,'chile':152,'china':156,'prc':156,'cn':156,'colombia':170,
  'congo':180,'costa rica':188,'croatia':191,'cuba':192,'czech republic':203,
  'czechia':203,'denmark':208,'dk':208,'dominican republic':214,
  'ecuador':218,'egypt':818,'eg':818,'ethiopia':231,'finland':246,'fi':246,
  'france':250,'fr':250,'germany':276,'de':276,'ghana':288,'greece':300,
  'gr':300,'hong kong':344,'hk':344,'hungary':348,'india':356,'in':356,
  'indonesia':360,'id':360,'iran':364,'iraq':368,'ireland':372,'israel':376,
  'il':376,'italy':380,'it':380,'jamaica':388,'japan':392,'jp':392,
  'jordan':400,'kenya':404,'south korea':410,'korea':410,'kr':410,
  'kuwait':414,'kw':414,'laos':418,'latvia':428,'libya':434,'malaysia':458,
  'my':458,'mexico':484,'mx':484,'morocco':504,'ma':504,'mozambique':508,
  'myanmar':104,'burma':104,'mm':104,'netherlands':528,'nl':528,'holland':528,
  'new zealand':554,'nz':554,'nicaragua':558,'nigeria':566,'ng':566,
  'norway':578,'no':578,'oman':512,'om':512,'pakistan':586,'pk':586,
  'panama':591,'pa':591,'peru':604,'philippines':608,'ph':608,'poland':616,
  'pl':616,'portugal':620,'pt':620,'qatar':634,'qa':634,'romania':642,
  'russia':643,'ru':643,'saudi arabia':682,'sa':682,'senegal':686,
  'singapore':702,'sg':702,'somalia':706,'south africa':710,'za':710,
  'spain':724,'es':724,'sri lanka':144,'lk':144,'sudan':729,'sd':729,
  'sweden':752,'se':752,'switzerland':756,'ch':756,'taiwan':158,'tw':158,
  'tanzania':834,'tz':834,'thailand':764,'th':764,'tunisia':788,'tn':788,
  'turkey':792,'tr':792,'ukraine':804,'ua':804,'united arab emirates':784,
  'uae':784,'ae':784,'united kingdom':826,'uk':826,'gb':826,'britain':826,
  'united states':840,'usa':840,'us':840,'america':840,'uruguay':858,
  'venezuela':862,'vietnam':704,'vn':704,'yemen':887,'ye':887,
  'zambia':894,'zimbabwe':716
};
(function() {
  var regions = ROUTE_DATA.highlight_regions || [];
  if (!regions.length) return;
  var codeColorMap = {};
  regions.forEach(function(r) {
    var key = (r.name || '').toLowerCase().trim();
    var code = COUNTRY_CODES_TRK[key];
    if (code !== undefined) codeColorMap[code] = {color: r.color, name: r.name};
  });
  if (!Object.keys(codeColorMap).length) return;
  fetch('https://cdn.jsdelivr.net/npm/world-atlas@2/countries-110m.json')
    .then(function(res) { return res.json(); })
    .then(function(world) {
      if (typeof topojson === 'undefined') return;
      var countries = topojson.feature(world, world.objects.countries);
      L.geoJSON(countries, {
        style: function(feature) {
          var match = codeColorMap[+feature.id];
          if (match) {
            var m = (match.color||'').match(/rgba?\\((\\d+),\\s*(\\d+),\\s*(\\d+)(?:,\\s*([\\d.]+))?\\)/);
            var hex = match.color;
            var alpha = 0.30;
            if (m) {
              hex = '#' + [+m[1],+m[2],+m[3]].map(function(v){return v.toString(16).padStart(2,'0');}).join('');
              alpha = m[4] !== undefined ? +m[4] : 0.30;
            }
            return {color: hex, weight: 1.5, opacity: 0.8, fillColor: hex, fillOpacity: alpha};
          }
          return {fillOpacity: 0, opacity: 0, weight: 0};
        },
        onEachFeature: function(feature, layer) {
          var match = codeColorMap[+feature.id];
          if (match) layer.bindTooltip('<b>' + match.name + '</b>', {sticky: true});
        }
      }).addTo(map);
    }).catch(function() {});
})();

// ── Arrow line helper (returns polyline for highlight control) ──────────────
function drawArrowLine(fromPt, toPt, color, laneOffset, tooltipHtml) {
  laneOffset = laneOffset || 0;
  var WSIZ = 1024;
  function merc(lat, lon) {
    var s = Math.sin(lat * Math.PI / 180);
    return { x: (lon + 180) / 360 * WSIZ, y: (0.5 - Math.log((1 + s) / (1 - s)) / (4 * Math.PI)) * WSIZ };
  }
  function unmerc(px, py) {
    var n = Math.PI - 2 * Math.PI * py / WSIZ;
    return { lat: 180 / Math.PI * Math.atan(0.5 * (Math.exp(n) - Math.exp(-n))), lng: px / WSIZ * 360 - 180 };
  }
  var p1 = merc(fromPt[0], fromPt[1]), p2 = merc(toPt[0], toPt[1]);
  var dx = p2.x - p1.x, dy = p2.y - p1.y;
  var pixChord = Math.sqrt(dx*dx + dy*dy) || 1;
  var STRAIGHT_THRESHOLD = 10;
  if (pixChord < STRAIGHT_THRESHOLD) {
    var _sl = L.polyline([fromPt, toPt], {color: color, weight: 3, opacity: 0.85}).addTo(map);
    if (tooltipHtml) _sl.bindTooltip(tooltipHtml, {sticky: true, className: 'route-tooltip'});
    var mLat = (fromPt[0]+toPt[0])/2, mLon = (fromPt[1]+toPt[1])/2;
    var mAngle = Math.atan2(toPt[1]-fromPt[1], toPt[0]-fromPt[0]) * 180/Math.PI;
    var mSvg = '<svg width="18" height="18" viewBox="-9 -9 18 18" xmlns="http://www.w3.org/2000/svg">'
      + '<polygon points="0,-7 5,3 0,0 -5,3" fill="' + color + '" opacity="0.95"'
      + ' transform="rotate(' + mAngle.toFixed(1) + ')"/></svg>';
    L.marker([mLat, mLon], {icon: L.divIcon({html: mSvg, className:'', iconSize:[18,18], iconAnchor:[9,9]}), interactive:false, zIndexOffset:100}).addTo(map);
    return _sl;
  }
  var perpX = -dy/pixChord, perpY = dx/pixChord;
  // Adaptive curve depth: flatter for nearby stops, gentler arc for distant ones
  var curveFactor = pixChord < 80 ? 0.04 : pixChord < 250 ? 0.05 : 0.09;
  var pxOffset = pixChord * curveFactor;
  var ctrl = unmerc(p1.x + dx/2 + perpX*laneOffset, p1.y + dy/2 - pxOffset + perpY*laneOffset);
  var ctrlLat = Math.min(85, Math.max(-85, ctrl.lat)), ctrlLon = ctrl.lng;
  var N = 50, curvePts = [];
  for (var i = 0; i <= N; i++) {
    var t = i/N, u = 1-t;
    curvePts.push([u*u*fromPt[0]+2*u*t*ctrlLat+t*t*toPt[0], u*u*fromPt[1]+2*u*t*ctrlLon+t*t*toPt[1]]);
  }
  var _cl = L.polyline(curvePts, {color: color, weight: 3, opacity: 0.85}).addTo(map);
  if (tooltipHtml) _cl.bindTooltip(tooltipHtml, {sticky: true, className: 'route-tooltip'});
  function bzPt(t) { var u=1-t; return [u*u*fromPt[0]+2*u*t*ctrlLat+t*t*toPt[0], u*u*fromPt[1]+2*u*t*ctrlLon+t*t*toPt[1]]; }
  var arrPt = bzPt(0.75), pa = bzPt(0.73), pb = bzPt(0.77);
  var angle = Math.atan2(pb[1]-pa[1], pb[0]-pa[0]) * 180/Math.PI;
  var svg = '<svg width="18" height="18" viewBox="-9 -9 18 18" xmlns="http://www.w3.org/2000/svg">'
    + '<polygon points="0,-7 5,3 0,0 -5,3" fill="' + color + '" opacity="0.95"'
    + ' transform="rotate(' + angle.toFixed(1) + ')"/></svg>';
  L.marker([arrPt[0], arrPt[1]], {icon: L.divIcon({html:svg, className:'', iconSize:[18,18], iconAnchor:[9,9]}), interactive:false, zIndexOffset:100}).addTo(map);
  return _cl;
}

// \u2500\u2500 Build popup HTML for a location \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
function buildPopup(loc) {
  // Count total lines: 1 (header) + 1 per track header + 1 per event
  var lineCount = 1;
  loc.tracks.forEach(function(t) { lineCount += 1 + t.events.length; });
  // Show scrollbar only when content exceeds 10 lines (~19px each + header padding ≈ 210px)
  var maxH = lineCount > 10 ? '210px' : 'none';
  var overflow = lineCount > 10 ? 'auto' : 'visible';
  var html = '<div style="max-height:' + maxH + ';overflow-y:' + overflow + ';min-width:220px;font-size:12px;line-height:1.6;">';
  html += '<div style="font-weight:bold;text-decoration:underline;text-align:center;padding-bottom:4px;white-space:nowrap;">' + loc.name + '</div>';
  loc.tracks.forEach(function(t) {
    html += '<div style="margin-top:4px;font-weight:bold;">' + t.trk + ' <span style="color:#555;font-weight:normal;">(' + t.routeType + '):</span></div>';
    t.events.forEach(function(ev) {
      html += '<div style="padding-left:10px;color:#333;">&nbsp;|&mdash;&mdash; ' + ev + '</div>';
    });
  });
  html += '</div>';
  return html;
}

// \u2500\u2500 Build tooltip HTML for a route line \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
function buildLineTooltip(route, fromLoc, toLoc) {
  var trk = route.trk;
  var noc = route.noOfContainers || 0;
  var containerLabel = noc === 1 ? '1 Container' : noc + ' Containers';
  function eventsFor(loc, trkName) {
    var t = null;
    for (var i = 0; i < loc.tracks.length; i++) { if (loc.tracks[i].trk === trkName) { t = loc.tracks[i]; break; } }
    return t ? t.events : [];
  }
  var html = '<div style="font-size:11px;line-height:1.7;white-space:nowrap;padding:4px 8px;min-width:180px;">';
  // Header: TrackNumber (N Containers) — bold, centred
  html += '<div style="font-weight:bold;text-align:center;margin-bottom:4px;border-bottom:1px solid #ddd;padding-bottom:3px;">'
       + trk + ' (' + containerLabel + ')</div>';
  // From section
  html += '<div><b>From:</b> ' + fromLoc.name + '</div>';
  eventsFor(fromLoc, trk).forEach(function(ev) {
    html += '<div style="padding-left:14px;">&nbsp;|-- ' + ev + '</div>';
  });
  // To section
  html += '<div style="margin-top:3px;"><b>To:</b> ' + toLoc.name + '</div>';
  eventsFor(toLoc, trk).forEach(function(ev) {
    html += '<div style="padding-left:14px;">&nbsp;|-- ' + ev + '</div>';
  });
  html += '</div>';
  return html;
}

// \u2500\u2500 Draw route lines \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var routeLayers = {};   // trk -> { lines: [polyline, ...], markerIdxs: [locIdx, ...] }
var lineGroups  = {};   // "fi-ti" -> offset count (for lane separation)
ROUTE_DATA.routes.forEach(function(route) {
  routeLayers[route.trk] = {lines: [], markerIdxs: []};
  for (var i = 0; i < route.stops.length - 1; i++) {
    var fi = route.stops[i], ti = route.stops[i+1];
    var key = Math.min(fi,ti) + '-' + Math.max(fi,ti);
    var offset = (lineGroups[key] || 0) * 2;
    lineGroups[key] = (lineGroups[key] || 0) + 1;
    var fromLoc = ROUTE_DATA.locations[fi], toLoc = ROUTE_DATA.locations[ti];
    var tip = buildLineTooltip(route, fromLoc, toLoc);
    var line = drawArrowLine([fromLoc.lat, fromLoc.lon], [toLoc.lat, toLoc.lon], route.color, offset, tip);
    if (line) routeLayers[route.trk].lines.push(line);
  }
  route.stops.forEach(function(idx) { routeLayers[route.trk].markerIdxs.push(idx); });
});

// \u2500\u2500 Draw location dots \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var locationMarkers = [];
ROUTE_DATA.locations.forEach(function(loc, idx) {
  var marker = L.circleMarker([loc.lat, loc.lon], {
    radius: 7, color: '#ffffff', weight: 2,
    fillColor: '#2c3e50', fillOpacity: 0.9
  }).addTo(map);
  marker.bindPopup(buildPopup(loc), {maxWidth: 420});
  marker.bindTooltip(loc.name, {permanent: true, className: 'stop-label', direction: 'top', offset: [0, -8]});
  locationMarkers.push(marker);
});

// \u2500\u2500 Fit bounds \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var allCoords = ROUTE_DATA.locations.map(function(l) { return [l.lat, l.lon]; });
if (allCoords.length > 0) {
  try { map.fitBounds(L.latLngBounds(allCoords), {padding: [60, 60], maxZoom: 10}); } catch(e) {}
}

// \u2500\u2500 Legend with hover highlight \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var legendCtrl = L.control({position: 'bottomright'});
legendCtrl.onAdd = function() {
  var div = L.DomUtil.create('div', 'legend');
  var h = '<h4>Legend</h4>';
  ROUTE_DATA.routes.forEach(function(route) {
    h += '<div class="legend-item" data-trk="' + route.trk + '">'
      + '<div class="leg-line" style="background:' + route.color + '"></div>'
      + '<div class="leg-dot" style="background:' + route.color + '"></div>'
      + '<span>' + route.trk + '</span></div>';
  });
  (ROUTE_DATA.highlight_regions || []).forEach(function(r) {
    var m = (r.color||'').match(/rgba?\\((\\d+),\\s*(\\d+),\\s*(\\d+)/);
    var hex = r.color || '#ffa500';
    if (m) hex = '#' + [+m[1],+m[2],+m[3]].map(function(v){return v.toString(16).padStart(2,'0');}).join('');
    h += '<div class="legend-item"><div class="leg-zone" style="background:' + hex
       + ';opacity:0.55;border:2px solid ' + hex + '"></div><span>' + r.name + '</span></div>';
  });
  h += '<div style="margin-top:8px;border-top:1px solid #ddd;padding-top:6px;">'
     + '<div class="legend-item" id="wz-toggle" title="Click to toggle war zone overlay" style="cursor:pointer;">'
     + '<div style="width:22px;height:11px;border:2px dashed #c0392b;background:rgba(231,76,60,0.18);flex-shrink:0;border-radius:2px;"></div>'
     + '<span style="color:#c0392b;font-weight:600;">&#9888; War Zones</span></div></div>';
  div.innerHTML = h;
  div.querySelectorAll('.legend-item').forEach(function(item) {
    var trk = item.getAttribute('data-trk');
    item.addEventListener('mouseenter', function() {
      // Dim all, then highlight hovered
      ROUTE_DATA.routes.forEach(function(r) {
        var rl = routeLayers[r.trk];
        if (!rl) return;
        var dimmed = (r.trk !== trk);
        rl.lines.forEach(function(l) { l.setStyle({opacity: dimmed ? 0.12 : 1.0, weight: dimmed ? 2 : 5}); });
        rl.markerIdxs.forEach(function(mi) {
          if (locationMarkers[mi]) locationMarkers[mi].setStyle({fillOpacity: dimmed ? 0.12 : 1.0, opacity: dimmed ? 0.2 : 1.0});
        });
      });
    });
    item.addEventListener('mouseleave', function() {
      ROUTE_DATA.routes.forEach(function(r) {
        var rl = routeLayers[r.trk];
        if (!rl) return;
        rl.lines.forEach(function(l) { l.setStyle({opacity: 0.85, weight: 3}); });
        rl.markerIdxs.forEach(function(mi) {
          if (locationMarkers[mi]) locationMarkers[mi].setStyle({fillOpacity: 0.9, opacity: 1.0});
        });
      });
    });
  });
  // War zone toggle
  var wzToggle = div.querySelector('#wz-toggle');
  if (wzToggle) {
    wzToggle.addEventListener('click', function() {
      warZonesVisible = !warZonesVisible;
      warZoneLayers.forEach(function(l) {
        if (warZonesVisible) { map.addLayer(l); } else { map.removeLayer(l); }
      });
      wzToggle.style.opacity = warZonesVisible ? '1' : '0.4';
    });
  }
  L.DomEvent.disableScrollPropagation(div);
  return div;
};
legendCtrl.addTo(map);

} catch(e) {
  var d = document.getElementById('map-error');
  if (d) { d.style.display='block'; d.innerHTML='<b>Map Error:</b> ' + e.message; }
}
</script>
</body>
</html>
"""

    route_data_json = _json.dumps(data, ensure_ascii=False)
    escaped_title = _html.escape(title)
    html_content = (
        _TRACKING_ROUTE_HTML
        .replace("|||TITLE|||", escaped_title)
        .replace("|||ROUTE_DATA_JSON|||", route_data_json)
    )
    full_path = f"{file_store_path}/{file_id}_{title_slug}.html"
    with open(full_path, "w", encoding="utf-8") as fh:
        fh.write(html_content)
    return _file_meta(file_id, f"{title_slug}.html", "text/html", full_path)


def _plot_location_bubble_map(
    title: str,
    locations: list[dict],
    value_label: str,
    default_color: str,
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Generate a Leaflet.js bubble/pin map for a list of locations.

    Each location is shown as a circle marker. If `value` is present the
    circles are proportionally sized. Tooltip shows name + optional value.
    """
    import json as _json
    import html as _html

    loc_json      = _json.dumps(locations, ensure_ascii=False)
    escaped_title = _html.escape(title)
    escaped_vlabel = _html.escape(value_label)

    _BUBBLE_MAP_HTML = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>|||TITLE|||</title>
<script>
window.onerror=function(m,s,l,c,e){
  var el=document.getElementById('map-error');
  if(el){el.style.display='block';el.innerHTML='<b>JS Error:</b> '+m+' (line '+l+')';}
  return false;
};
</script>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>
  html, body, #map { margin:0; padding:0; width:100%; height:100vh; overflow:hidden; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }
  .map-title {
    position:absolute; top:12px; left:50%; transform:translateX(-50%);
    z-index:1000; background:rgba(255,255,255,0.95);
    padding:7px 20px; border-radius:8px;
    font-size:15px; font-weight:700; color:#1F4788;
    box-shadow:0 2px 10px rgba(0,0,0,0.18);
    white-space:nowrap; pointer-events:none;
  }
  .legend {
    background:rgba(255,255,255,0.97); padding:10px 14px;
    border-radius:8px; box-shadow:0 2px 10px rgba(0,0,0,0.18);
    font-size:12px; line-height:1.7; min-width:130px;
  }
  .legend h4 { margin:0 0 6px 0; font-size:13px; color:#1F4788; }
  .legend-item { display:flex; align-items:center; gap:7px; }
  .leg-bubble { border-radius:50%; border:2px solid rgba(255,255,255,0.9); flex-shrink:0; }
  .bubble-tooltip { background:rgba(255,255,255,0.97); border:1px solid #ccc; border-radius:5px; box-shadow:0 2px 8px rgba(0,0,0,0.15); }
  .stop-label {
    background:transparent !important; border:none !important;
    box-shadow:none !important; font-size:10px; font-weight:600;
    color:#1F4788; white-space:nowrap; pointer-events:none;
  }
  #map-error {
    display:none; position:absolute; top:60px; left:50%; transform:translateX(-50%);
    z-index:9999; background:#fff3cd; border:1px solid #ffc107;
    padding:10px 16px; border-radius:6px; font-size:13px;
  }
</style>
</head>
<body>
<div id="map"></div>
<div class="map-title">|||TITLE|||</div>
<div id="map-error"></div>
<script>
try {
var LOCS        = |||LOCS_JSON|||;
var VALUE_LABEL = "|||VALUE_LABEL|||";
var DEF_COLOR   = "|||DEF_COLOR|||";

var map = L.map('map', {preferCanvas: true}).setView([20, 10], 2);
L.tileLayer('https://{s}.basemaps.cartocdn.com/rastertiles/voyager/{z}/{x}/{y}{r}.png', {
  attribution: '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> &copy; <a href="https://carto.com/">CARTO</a>',
  subdomains: 'abcd', maxZoom: 19
}).addTo(map);
[100,300,600,1200].forEach(function(ms){setTimeout(function(){map.invalidateSize(false);},ms);});
window.addEventListener('resize', function(){map.invalidateSize(false);});

// ── Compute value range for proportional sizing ───────────────────────────
var hasValues = LOCS.some(function(l){return l.value !== undefined && l.value !== null;});
var maxVal = 0, minVal = Infinity;
if (hasValues) {
  LOCS.forEach(function(l) {
    var v = +l.value || 0;
    if (v > maxVal) maxVal = v;
    if (v < minVal) minVal = v;
  });
  if (!isFinite(minVal)) minVal = 0;
  if (maxVal === 0) maxVal = 1;
}

var MIN_R = 8, MAX_R = 36;
function radius(v) {
  if (!hasValues) return 10;
  var t = maxVal > minVal ? (v - minVal) / (maxVal - minVal) : 1;
  return MIN_R + t * (MAX_R - MIN_R);
}
function fmtNum(n) {
  return n >= 1000 ? (n/1000).toFixed(1).replace(/\\.0$/,'') + 'k' : String(n);
}

// ── Draw bubbles ──────────────────────────────────────────────────────────
var allCoords = [];
LOCS.forEach(function(loc) {
  var lat = +loc.lat, lon = +loc.lon;
  if (isNaN(lat) || isNaN(lon)) return;
  allCoords.push([lat, lon]);
  var v     = (loc.value !== undefined && loc.value !== null) ? +loc.value : null;
  var r     = radius(v !== null ? v : 0);
  var col   = loc.color || DEF_COLOR;
  var circle = L.circleMarker([lat, lon], {
    radius: r,
    color: '#ffffff', weight: 2,
    fillColor: col, fillOpacity: 0.80
  }).addTo(map);

  // Tooltip
  var tip = '<div style="font-size:12px;line-height:1.6;padding:2px 4px;">'
    + '<b>' + (loc.name || '') + '</b>';
  if (v !== null) tip += '<br>' + VALUE_LABEL + ': <b>' + (v.toLocaleString ? v.toLocaleString() : v) + '</b>';
  if (loc.label) tip += '<br><span style="color:#555">' + loc.label + '</span>';
  tip += '</div>';
  circle.bindTooltip(tip, {sticky: true, className: 'bubble-tooltip'});

  // Permanent label
  L.tooltip({permanent: true, className: 'stop-label', direction: 'top', offset: [0, -(r+4)]})
    .setContent(loc.name || '')
    .setLatLng([lat, lon])
    .addTo(map);
});

// ── Fit bounds ────────────────────────────────────────────────────────────
if (allCoords.length > 0) {
  try { map.fitBounds(L.latLngBounds(allCoords), {padding: [60, 60], maxZoom: 10}); } catch(e) {}
}

// ── Legend ────────────────────────────────────────────────────────────────
if (hasValues) {
  var legendCtrl = L.control({position: 'bottomright'});
  legendCtrl.onAdd = function() {
    var div = L.DomUtil.create('div', 'legend');
    var h = '<h4>' + VALUE_LABEL + '</h4>';
    var steps = [
      {label: fmtNum(maxVal), r: MAX_R},
      {label: fmtNum(Math.round((maxVal + minVal) / 2)), r: Math.round((MIN_R + MAX_R) / 2)},
      {label: fmtNum(minVal), r: MIN_R}
    ];
    steps.forEach(function(s) {
      h += '<div class="legend-item" style="margin-bottom:4px;">'
        + '<div class="leg-bubble" style="width:' + (s.r*2) + 'px;height:' + (s.r*2) + 'px;background:' + DEF_COLOR + ';opacity:0.8;"></div>'
        + '<span>' + s.label + '</span></div>';
    });
    div.innerHTML = h;
    L.DomEvent.disableScrollPropagation(div);
    return div;
  };
  legendCtrl.addTo(map);
}

} catch(e) {
  var d = document.getElementById('map-error');
  if (d) { d.style.display='block'; d.innerHTML='<b>Map Error:</b> ' + e.message; }
}
</script>
</body>
</html>
"""

    html_content = (
        _BUBBLE_MAP_HTML
        .replace("|||TITLE|||",      escaped_title)
        .replace("|||LOCS_JSON|||",  loc_json)
        .replace("|||VALUE_LABEL|||", escaped_vlabel)
        .replace("|||DEF_COLOR|||",  default_color)
    )
    full_path = f"{file_store_path}/{file_id}_{title_slug}.html"
    with open(full_path, "w", encoding="utf-8") as fh:
        fh.write(html_content)
    return _file_meta(file_id, f"{title_slug}.html", "text/html", full_path)


def _plot_choropleth_map(
    title: str,
    data: list[dict],
    color: str,
    value_label: str,
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Generate a Leaflet.js choropleth world map shading countries by value.

    Args:
        data: list of {country, value, label?} dicts
        color: base hue ('blue','red','green','orange','purple')
        value_label: tooltip label for the numeric value
    """
    import json as _json
    import html as _html

    # ── Base RGB for each hue ────────────────────────────────────────────
    HUES = {
        "blue":   (31,  71, 136),
        "red":    (180,  20,  20),
        "green":  ( 20, 120,  60),
        "orange": (200,  90,   0),
        "purple": ( 90,  20, 150),
    }
    base_r, base_g, base_b = HUES.get(color, HUES["blue"])

    data_json    = _json.dumps(data,        ensure_ascii=False)
    escaped_title = _html.escape(title)
    escaped_vlabel = _html.escape(value_label)

    _CHOROPLETH_HTML = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>|||TITLE|||</title>
<script>
window.onerror=function(m,s,l,c,e){
  var el=document.getElementById('map-error');
  if(el){el.style.display='block';el.innerHTML='<b>JS Error:</b> '+m+' (line '+l+')';}
  return false;
};
</script>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script src="https://cdn.jsdelivr.net/npm/topojson-client@3/dist/topojson-client.min.js"></script>
<style>
  html, body, #map { margin:0; padding:0; width:100%; height:100vh; overflow:hidden; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }
  .map-title {
    position:absolute; top:12px; left:50%; transform:translateX(-50%);
    z-index:1000; background:rgba(255,255,255,0.95);
    padding:7px 20px; border-radius:8px;
    font-size:15px; font-weight:700; color:#1F4788;
    box-shadow:0 2px 10px rgba(0,0,0,0.18);
    white-space:nowrap; pointer-events:none;
  }
  .legend {
    background:rgba(255,255,255,0.97); padding:10px 14px;
    border-radius:8px; box-shadow:0 2px 10px rgba(0,0,0,0.18);
    font-size:12px; line-height:1.7; min-width:150px;
  }
  .legend h4 { margin:0 0 6px 0; font-size:13px; color:#1F4788; }
  .legend-gradient {
    width:120px; height:12px; border-radius:3px;
    background: linear-gradient(to right, |||GRAD_LIGHT|||, |||GRAD_DARK|||);
    margin-bottom:4px;
  }
  .legend-labels { display:flex; justify-content:space-between; font-size:10px; color:#555; }
  .route-tooltip { background:rgba(255,255,255,0.97); border:1px solid #ccc; border-radius:5px; box-shadow:0 2px 8px rgba(0,0,0,0.15); }
  #map-error {
    display:none; position:absolute; top:60px; left:50%; transform:translateX(-50%);
    z-index:9999; background:#fff3cd; border:1px solid #ffc107;
    padding:10px 16px; border-radius:6px; font-size:13px;
  }
</style>
</head>
<body>
<div id="map"></div>
<div class="map-title">|||TITLE|||</div>
<div id="map-error"></div>
<script>
try {
var CHORO_DATA  = |||DATA_JSON|||;
var VALUE_LABEL = "|||VALUE_LABEL|||";
var BASE_RGB    = [|||BASE_R|||, |||BASE_G|||, |||BASE_B|||];

// ── Country name → ISO numeric code ───────────────────────────────────────
var COUNTRY_CODES = {
  'afghanistan':4,'albania':8,'algeria':12,'angola':24,'argentina':32,
  'australia':36,'austria':40,'bahrain':48,'bangladesh':50,'belgium':56,
  'bolivia':68,'brazil':76,'bulgaria':100,'cambodia':116,'cameroon':120,
  'canada':124,'chile':152,'china':156,'prc':156,'cn':156,'colombia':170,
  'congo':180,'costa rica':188,'croatia':191,'cuba':192,'czech republic':203,
  'czechia':203,'denmark':208,'dk':208,'dominican republic':214,
  'ecuador':218,'egypt':818,'eg':818,'ethiopia':231,'finland':246,'fi':246,
  'france':250,'fr':250,'germany':276,'de':276,'ghana':288,'greece':300,
  'gr':300,'hong kong':344,'hk':344,'hungary':348,'india':356,'in':356,
  'indonesia':360,'id':360,'iran':364,'iraq':368,'ireland':372,'israel':376,
  'il':376,'italy':380,'it':380,'jamaica':388,'japan':392,'jp':392,
  'jordan':400,'kenya':404,'south korea':410,'korea':410,'kr':410,
  'kuwait':414,'kw':414,'laos':418,'latvia':428,'libya':434,'malaysia':458,
  'my':458,'mexico':484,'mx':484,'morocco':504,'ma':504,'mozambique':508,
  'myanmar':104,'burma':104,'mm':104,'netherlands':528,'nl':528,'holland':528,
  'new zealand':554,'nz':554,'nicaragua':558,'nigeria':566,'ng':566,
  'norway':578,'no':578,'oman':512,'om':512,'pakistan':586,'pk':586,
  'panama':591,'pa':591,'peru':604,'philippines':608,'ph':608,'poland':616,
  'pl':616,'portugal':620,'pt':620,'qatar':634,'qa':634,'romania':642,
  'russia':643,'ru':643,'saudi arabia':682,'sa':682,'senegal':686,
  'singapore':702,'sg':702,'somalia':706,'south africa':710,'za':710,
  'spain':724,'es':724,'sri lanka':144,'lk':144,'sudan':729,'sd':729,
  'sweden':752,'se':752,'switzerland':756,'ch':756,'taiwan':158,'tw':158,
  'tanzania':834,'tz':834,'thailand':764,'th':764,'tunisia':788,'tn':788,
  'turkey':792,'tr':792,'ukraine':804,'ua':804,'united arab emirates':784,
  'uae':784,'ae':784,'united kingdom':826,'uk':826,'gb':826,'britain':826,
  'united states':840,'usa':840,'us':840,'america':840,'uruguay':858,
  'venezuela':862,'vietnam':704,'vn':704,'yemen':887,'ye':887,
  'zambia':894,'zimbabwe':716
};

// ── Build code → {value, label, country} map ──────────────────────────────
var codeMap = {};
var maxVal  = 0, minVal = Infinity;
CHORO_DATA.forEach(function(d) {
  var key  = (d.country || '').toLowerCase().trim();
  var code = COUNTRY_CODES[key];
  if (code === undefined) return;
  var v = +d.value || 0;
  if (v > maxVal) maxVal = v;
  if (v < minVal) minVal = v;
  codeMap[code] = { value: v, label: d.label || d.country, country: d.country };
});
if (maxVal === 0) maxVal = 1;
if (!isFinite(minVal)) minVal = 0;

function valueToColor(v) {
  // Log scale for better visual distribution across wide value ranges
  var logMin = minVal > 0 ? Math.log(minVal) : 0;
  var logMax = Math.log(maxVal + 1);
  var logV   = Math.log(v + 1);
  var t = logMax > logMin ? (logV - logMin) / (logMax - logMin) : 1;
  t = Math.max(0.08, Math.min(1.0, t));  // keep min visible at 8%
  // Light (t=0) → full base color (t=1); background blends to white
  var r = Math.round(BASE_RGB[0] * t + 220 * (1 - t));
  var g = Math.round(BASE_RGB[1] * t + 230 * (1 - t));
  var b = Math.round(BASE_RGB[2] * t + 245 * (1 - t));
  return { hex: '#' + [r,g,b].map(function(x){return Math.min(255,x).toString(16).padStart(2,'0');}).join(''), alpha: 0.25 + t * 0.55 };
}

function fmtNum(n) {
  return n >= 1000 ? (n/1000).toFixed(1).replace(/\\.0$/,'') + 'k' : String(n);
}

// ── Init map ──────────────────────────────────────────────────────────────
var map = L.map('map', {preferCanvas: true}).setView([20, 10], 2);
L.tileLayer('https://{s}.basemaps.cartocdn.com/rastertiles/voyager/{z}/{x}/{y}{r}.png', {
  attribution: '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> &copy; <a href="https://carto.com/">CARTO</a>',
  subdomains: 'abcd', maxZoom: 19
}).addTo(map);
[100,300,600,1200].forEach(function(ms){setTimeout(function(){map.invalidateSize(false);},ms);});
window.addEventListener('resize', function(){map.invalidateSize(false);});

// ── Render choropleth ─────────────────────────────────────────────────────
fetch('https://cdn.jsdelivr.net/npm/world-atlas@2/countries-110m.json')
  .then(function(r){return r.json();})
  .then(function(world){
    if (typeof topojson === 'undefined') { throw new Error('topojson not loaded'); }
    var countries = topojson.feature(world, world.objects.countries);
    var geoLayer = L.geoJSON(countries, {
      style: function(feature) {
        var match = codeMap[+feature.id];
        if (match) {
          var c = valueToColor(match.value);
          return { color: c.hex, weight: 0.8, opacity: 0.7, fillColor: c.hex, fillOpacity: c.alpha };
        }
        return { color: '#cccccc', weight: 0.5, opacity: 0.4, fillColor: '#f0f0f0', fillOpacity: 0.3 };
      },
      onEachFeature: function(feature, layer) {
        var match = codeMap[+feature.id];
        if (match) {
          layer.bindTooltip(
            '<div style="font-size:12px;line-height:1.6;padding:2px 4px;">'
            + '<b>' + match.label + '</b><br>'
            + VALUE_LABEL + ': <b>' + match.value.toLocaleString() + '</b>'
            + '</div>',
            { sticky: true, className: 'route-tooltip' }
          );
          layer.on('mouseover', function(e) {
            layer.setStyle({ weight: 2, opacity: 1.0, fillOpacity: Math.min(1, (valueToColor(match.value).alpha + 0.2)) });
          });
          layer.on('mouseout', function(e) {
            geoLayer.resetStyle(layer);
          });
        }
      }
    }).addTo(map);
  }).catch(function(e){
    var el = document.getElementById('map-error');
    if (el) { el.style.display='block'; el.innerHTML='<b>Map Error:</b> ' + e.message; }
  });

// ── Legend ────────────────────────────────────────────────────────────────
var legendCtrl = L.control({position: 'bottomright'});
legendCtrl.onAdd = function() {
  var div = L.DomUtil.create('div', 'legend');
  div.innerHTML = '<h4>' + VALUE_LABEL + '</h4>'
    + '<div class="legend-gradient"></div>'
    + '<div class="legend-labels"><span>' + fmtNum(minVal) + '</span><span>' + fmtNum(maxVal) + '</span></div>';
  L.DomEvent.disableScrollPropagation(div);
  return div;
};
legendCtrl.addTo(map);

} catch(e) {
  var d = document.getElementById('map-error');
  if (d) { d.style.display='block'; d.innerHTML='<b>Map Error:</b> ' + e.message; }
}
</script>
</body>
</html>
"""

    # Compute gradient light/dark hex for CSS
    def _to_hex(r: int, g: int, b: int) -> str:
        return "#{:02x}{:02x}{:02x}".format(r, g, b)

    grad_light = _to_hex(
        int(base_r * 0.08 + 220 * 0.92),
        int(base_g * 0.08 + 230 * 0.92),
        int(base_b * 0.08 + 245 * 0.92),
    )
    grad_dark  = _to_hex(base_r, base_g, base_b)

    html_content = (
        _CHOROPLETH_HTML
        .replace("|||TITLE|||",       escaped_title)
        .replace("|||DATA_JSON|||",   data_json)
        .replace("|||VALUE_LABEL|||", escaped_vlabel)
        .replace("|||BASE_R|||",      str(base_r))
        .replace("|||BASE_G|||",      str(base_g))
        .replace("|||BASE_B|||",      str(base_b))
        .replace("|||GRAD_LIGHT|||",  grad_light)
        .replace("|||GRAD_DARK|||",   grad_dark)
    )
    full_path = f"{file_store_path}/{file_id}_{title_slug}.html"
    with open(full_path, "w", encoding="utf-8") as fh:
        fh.write(html_content)
    return _file_meta(file_id, f"{title_slug}.html", "text/html", full_path)


def _plot_container_route_map(
    title: str,
    data: dict[str, Any],
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Generate a Leaflet.js container route map as a self-contained HTML file.

    Data format expected:
      locations: [{name, lat, lon, containers:[{key, events:[str]}]}]
      routes:    [{key, trk, color, stops:[loc_idx], vessels:[str]}]

    Features:
      - One dot per unique lat/lon (shared across containers)
      - Popup: LocationName header + sorted containers + events + scrollbar if >10 lines
      - Route lines per container (no cross-container connections)
      - Tooltip: ContainerKey – Vessel at From, From/To with last/first event
      - Per-TrackNumber base color, light→dark gradient per container
      - Legend per container with hover highlight
      - War zone overlays with toggle
      - CartoDB Voyager tiles (English labels)
    """
    import json as _json
    import html as _html

    _CONTAINER_ROUTE_HTML = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>|||TITLE|||</title>
<script>
window.onerror=function(m,s,l,c,e){
  var el=document.getElementById('map-error');
  if(el){el.style.display='block';el.innerHTML='<b>JS Error:</b> '+m+' (line '+l+')';}
  return false;
};
</script>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script src="https://cdn.jsdelivr.net/npm/topojson-client@3/dist/topojson-client.min.js"></script>
<style>
  html, body, #map { margin: 0; padding: 0; width: 100%; height: 100vh; overflow: hidden; }
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }
  .map-title {
    position: absolute; top: 12px; left: 50%; transform: translateX(-50%);
    z-index: 1000; background: rgba(255,255,255,0.95);
    padding: 7px 20px; border-radius: 8px;
    font-size: 15px; font-weight: 700; color: #1F4788;
    box-shadow: 0 2px 10px rgba(0,0,0,0.18);
    white-space: nowrap; pointer-events: none;
  }
  .legend {
    background: rgba(255,255,255,0.97); padding: 10px 14px;
    border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.18);
    font-size: 12px; line-height: 1.7; max-height: 340px; overflow-y: auto; min-width: 160px;
  }
  .legend h4 { margin: 0 0 6px 0; font-size: 13px; color: #1F4788; }
  .legend-item { display: flex; align-items: center; gap: 7px; padding: 1px 4px; border-radius: 3px; cursor: pointer; }
  .legend-item:hover { background: #f0f4ff; }
  .leg-swatch { width: 22px; height: 10px; border-radius: 3px; flex-shrink: 0; }
  .leg-zone { width: 22px; height: 11px; border-radius: 3px; flex-shrink: 0; }
  .route-tooltip { background: rgba(255,255,255,0.97); border: 1px solid #ccc; border-radius: 5px; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }
  .stop-label {
    background: transparent !important; border: none !important;
    box-shadow: none !important; font-size: 10px; font-weight: 600;
    color: #1F4788; white-space: nowrap;
  }
  #map-error {
    display: none; position: absolute; top: 60px; left: 50%; transform: translateX(-50%);
    z-index: 9999; background: #fff3cd; border: 1px solid #ffc107;
    padding: 10px 16px; border-radius: 6px; font-size: 13px;
  }
</style>
</head>
<body>
<div id="map"></div>
<div class="map-title">|||TITLE|||</div>
<div id="map-error"></div>
<script>
try {
var ROUTE_DATA = |||ROUTE_DATA_JSON|||;

var map = L.map('map', {preferCanvas: true}).setView([20, 0], 2);
// CartoDB Voyager — renders all labels in English
L.tileLayer('https://{s}.basemaps.cartocdn.com/rastertiles/voyager/{z}/{x}/{y}{r}.png', {
  attribution: '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors &copy; <a href="https://carto.com/attributions">CARTO</a>',
  subdomains: 'abcd',
  maxZoom: 20
}).addTo(map);

// \u2500\u2500 War Zone overlays \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var WAR_ZONES = [
  { name: "Red Sea / Gulf of Aden / Bab el-Mandeb (Houthi High Risk Zone)",
    coords: [[18.0,37.8],[18.0,41.0],[16.0,43.2],[13.5,43.5],[12.5,43.5],[12.0,46.0],[12.5,49.0],[13.5,51.5],[16.5,53.0],[13.5,57.0],[11.0,60.0],[6.0,50.0],[-1.5,41.5],[3.0,42.0],[7.0,41.5],[10.5,42.5],[11.5,43.2],[14.0,41.0],[16.5,39.5],[18.0,37.8]] },
  { name: "Gaza / Israel (Conflict Zone)",
    coords: [[29.2,33.8],[29.2,36.5],[33.5,36.5],[33.5,33.8]] },
  { name: "Ukraine / Black Sea (Conflict Zone)",
    coords: [[41.0,28.0],[43.5,28.0],[46.5,30.0],[47.0,32.0],[47.1,35.0],[46.5,38.0],[47.5,38.5],[47.0,39.5],[45.5,41.5],[43.5,41.5],[41.5,41.5],[41.0,40.0],[41.0,28.0]] },
  { name: "Sudan (Civil War Zone)",
    coords: [[12.0,22.5],[22.0,22.5],[22.0,37.5],[18.5,40.0],[15.5,38.5],[12.5,36.5],[10.0,33.0],[10.0,23.5],[12.0,22.5]] }
];
var warZoneLayers = WAR_ZONES.map(function(z) {
  var poly = L.polygon(z.coords, {
    color: '#c0392b', weight: 1.5, opacity: 0.85,
    fillColor: '#e74c3c', fillOpacity: 0.18, dashArray: '6,4', interactive: true
  }).addTo(map);
  poly.bindTooltip('<div style="font-weight:bold;color:#c0392b;font-size:12px;white-space:nowrap;">&#9888; ' + z.name + '</div>',
    {sticky: true, className: 'route-tooltip'});
  return poly;
});
var warZonesVisible = true;

// ── Country / region highlights (TopoJSON world atlas) ─────────────────────
var COUNTRY_CODES_CTR = {
  'afghanistan':4,'albania':8,'algeria':12,'angola':24,'argentina':32,
  'australia':36,'austria':40,'bahrain':48,'bangladesh':50,'belgium':56,
  'bolivia':68,'brazil':76,'bulgaria':100,'cambodia':116,'cameroon':120,
  'canada':124,'chile':152,'china':156,'prc':156,'cn':156,'colombia':170,
  'congo':180,'costa rica':188,'croatia':191,'cuba':192,'czech republic':203,
  'czechia':203,'denmark':208,'dk':208,'dominican republic':214,
  'ecuador':218,'egypt':818,'eg':818,'ethiopia':231,'finland':246,'fi':246,
  'france':250,'fr':250,'germany':276,'de':276,'ghana':288,'greece':300,
  'gr':300,'hong kong':344,'hk':344,'hungary':348,'india':356,'in':356,
  'indonesia':360,'id':360,'iran':364,'iraq':368,'ireland':372,'israel':376,
  'il':376,'italy':380,'it':380,'jamaica':388,'japan':392,'jp':392,
  'jordan':400,'kenya':404,'south korea':410,'korea':410,'kr':410,
  'kuwait':414,'kw':414,'laos':418,'latvia':428,'libya':434,'malaysia':458,
  'my':458,'mexico':484,'mx':484,'morocco':504,'ma':504,'mozambique':508,
  'myanmar':104,'burma':104,'mm':104,'netherlands':528,'nl':528,'holland':528,
  'new zealand':554,'nz':554,'nicaragua':558,'nigeria':566,'ng':566,
  'norway':578,'no':578,'oman':512,'om':512,'pakistan':586,'pk':586,
  'panama':591,'pa':591,'peru':604,'philippines':608,'ph':608,'poland':616,
  'pl':616,'portugal':620,'pt':620,'qatar':634,'qa':634,'romania':642,
  'russia':643,'ru':643,'saudi arabia':682,'sa':682,'senegal':686,
  'singapore':702,'sg':702,'somalia':706,'south africa':710,'za':710,
  'spain':724,'es':724,'sri lanka':144,'lk':144,'sudan':729,'sd':729,
  'sweden':752,'se':752,'switzerland':756,'ch':756,'taiwan':158,'tw':158,
  'tanzania':834,'tz':834,'thailand':764,'th':764,'tunisia':788,'tn':788,
  'turkey':792,'tr':792,'ukraine':804,'ua':804,'united arab emirates':784,
  'uae':784,'ae':784,'united kingdom':826,'uk':826,'gb':826,'britain':826,
  'united states':840,'usa':840,'us':840,'america':840,'uruguay':858,
  'venezuela':862,'vietnam':704,'vn':704,'yemen':887,'ye':887,
  'zambia':894,'zimbabwe':716
};
(function() {
  var regions = ROUTE_DATA.highlight_regions || [];
  if (!regions.length) return;
  var codeColorMap = {};
  regions.forEach(function(r) {
    var key = (r.name || '').toLowerCase().trim();
    var code = COUNTRY_CODES_CTR[key];
    if (code !== undefined) codeColorMap[code] = {color: r.color, name: r.name};
  });
  if (!Object.keys(codeColorMap).length) return;
  fetch('https://cdn.jsdelivr.net/npm/world-atlas@2/countries-110m.json')
    .then(function(res) { return res.json(); })
    .then(function(world) {
      if (typeof topojson === 'undefined') return;
      var countries = topojson.feature(world, world.objects.countries);
      L.geoJSON(countries, {
        style: function(feature) {
          var match = codeColorMap[+feature.id];
          if (match) {
            var m = (match.color||'').match(/rgba?\\((\\d+),\\s*(\\d+),\\s*(\\d+)(?:,\\s*([\\d.]+))?\\)/);
            var hex = match.color;
            var alpha = 0.30;
            if (m) {
              hex = '#' + [+m[1],+m[2],+m[3]].map(function(v){return v.toString(16).padStart(2,'0');}).join('');
              alpha = m[4] !== undefined ? +m[4] : 0.30;
            }
            return {color: hex, weight: 1.5, opacity: 0.8, fillColor: hex, fillOpacity: alpha};
          }
          return {fillOpacity: 0, opacity: 0, weight: 0};
        },
        onEachFeature: function(feature, layer) {
          var match = codeColorMap[+feature.id];
          if (match) layer.bindTooltip('<b>' + match.name + '</b>', {sticky: true});
        }
      }).addTo(map);
    }).catch(function() {});
})();

// ── Arrow curve line helper ─────────────────────────────────────────────────
function drawArrowLine(fromPt, toPt, color, laneOffset, tooltipHtml) {
  laneOffset = laneOffset || 0;
  var WSIZ = 1024;
  function merc(lat, lon) {
    var s = Math.sin(lat * Math.PI / 180);
    return { x: (lon + 180) / 360 * WSIZ, y: (0.5 - Math.log((1 + s) / (1 - s)) / (4 * Math.PI)) * WSIZ };
  }
  function unmerc(px, py) {
    var n = Math.PI - 2 * Math.PI * py / WSIZ;
    return { lat: 180 / Math.PI * Math.atan(0.5 * (Math.exp(n) - Math.exp(-n))), lng: px / WSIZ * 360 - 180 };
  }
  var p1 = merc(fromPt[0], fromPt[1]), p2 = merc(toPt[0], toPt[1]);
  var dx = p2.x - p1.x, dy = p2.y - p1.y;
  var pixChord = Math.sqrt(dx*dx + dy*dy) || 1;
  var STRAIGHT_THRESHOLD = 10;
  if (pixChord < STRAIGHT_THRESHOLD) {
    var _sl = L.polyline([fromPt, toPt], {color: color, weight: 1.8, opacity: 0.85}).addTo(map);
    if (tooltipHtml) _sl.bindTooltip(tooltipHtml, {sticky: true, className: 'route-tooltip'});
    var mLat = (fromPt[0]+toPt[0])/2, mLon = (fromPt[1]+toPt[1])/2;
    var mAngle = Math.atan2(toPt[1]-fromPt[1], toPt[0]-fromPt[0]) * 180/Math.PI;
    var mSvg = '<svg width="16" height="16" viewBox="-8 -8 16 16" xmlns="http://www.w3.org/2000/svg">'
      + '<polygon points="0,-6 5,3 0,0 -5,3" fill="' + color + '" opacity="0.95"'
      + ' transform="rotate(' + mAngle.toFixed(1) + ')"/></svg>';
    L.marker([mLat, mLon], {icon: L.divIcon({html: mSvg, className:'', iconSize:[16,16], iconAnchor:[8,8]}), interactive:false, zIndexOffset:100}).addTo(map);
    return _sl;
  }
  var perpX = -dy/pixChord, perpY = dx/pixChord;
  // Adaptive curve depth: flatter for nearby stops, gentler arc for distant ones
  var curveFactor = pixChord < 80 ? 0.04 : pixChord < 250 ? 0.05 : 0.09;
  var pxOffset = pixChord * curveFactor;
  var ctrl = unmerc(p1.x + dx/2 + perpX*laneOffset, p1.y + dy/2 - pxOffset + perpY*laneOffset);
  var ctrlLat = Math.min(85, Math.max(-85, ctrl.lat)), ctrlLon = ctrl.lng;
  var N = 50, curvePts = [];
  for (var i = 0; i <= N; i++) {
    var t = i/N, u = 1-t;
    curvePts.push([u*u*fromPt[0]+2*u*t*ctrlLat+t*t*toPt[0], u*u*fromPt[1]+2*u*t*ctrlLon+t*t*toPt[1]]);
  }
  var _cl = L.polyline(curvePts, {color: color, weight: 1.8, opacity: 0.85}).addTo(map);
  if (tooltipHtml) _cl.bindTooltip(tooltipHtml, {sticky: true, className: 'route-tooltip'});
  function bzPt(t) { var u=1-t; return [u*u*fromPt[0]+2*u*t*ctrlLat+t*t*toPt[0], u*u*fromPt[1]+2*u*t*ctrlLon+t*t*toPt[1]]; }
  var arrPt = bzPt(0.75), pa = bzPt(0.73), pb = bzPt(0.77);
  var angle = Math.atan2(pb[1]-pa[1], pb[0]-pa[0]) * 180/Math.PI;
  var svg = '<svg width="16" height="16" viewBox="-8 -8 16 16" xmlns="http://www.w3.org/2000/svg">'
    + '<polygon points="0,-6 5,3 0,0 -5,3" fill="' + color + '" opacity="0.95"'
    + ' transform="rotate(' + angle.toFixed(1) + ')"/></svg>';
  L.marker([arrPt[0], arrPt[1]], {icon: L.divIcon({html:svg, className:'', iconSize:[16,16], iconAnchor:[8,8]}), interactive:false, zIndexOffset:100}).addTo(map);
  return _cl;
}

// \u2500\u2500 Popup: LocationName header + sorted containers + events \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
function buildContainerPopup(loc) {
  // Count total lines: 1 (location header) + per container: 1 (name) + N (events)
  var lineCount = 1;
  loc.containers.forEach(function(c) { lineCount += 1 + c.events.length; });
  var maxH    = lineCount > 10 ? '210px' : 'none';
  var overflow = lineCount > 10 ? 'auto'  : 'visible';
  var html = '<div style="max-height:' + maxH + ';overflow-y:' + overflow + ';min-width:220px;font-size:12px;line-height:1.6;">';
  html += '<div style="font-weight:bold;text-decoration:underline;text-align:center;padding-bottom:4px;white-space:nowrap;">' + loc.name + '</div>';
  loc.containers.forEach(function(c) {
    html += '<div style="margin-top:4px;font-weight:600;">' + c.key + ':</div>';
    c.events.forEach(function(ev) {
      html += '<div style="padding-left:10px;color:#333;">&nbsp;|-- ' + ev + '</div>';
    });
  });
  html += '</div>';
  return html;
}

// \u2500\u2500 Tooltip for a route line segment \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
// Format: ContainerKey \u2013 Vessel at From
//         From: location name
//               |-- last event at From
//         To:   location name
//               |-- first event at To
function buildContainerLineTooltip(route, fromLoc, toLoc, vessel) {
  function eventsFor(loc, ckey) {
    for (var i = 0; i < loc.containers.length; i++) {
      if (loc.containers[i].key === ckey) return loc.containers[i].events;
    }
    return [];
  }
  var fromEvents = eventsFor(fromLoc, route.key);
  var toEvents   = eventsFor(toLoc,   route.key);
  var lastFrom   = fromEvents.length > 0 ? fromEvents[fromEvents.length - 1] : '';
  var firstTo    = toEvents.length   > 0 ? toEvents[0]                       : '';

  var html = '<div style="font-size:11px;line-height:1.7;white-space:nowrap;padding:4px 8px;min-width:200px;">';
  // Header: ContainerKey \u2013 Vessel (bold, bottom-bordered)
  var header = route.key;
  if (vessel) header += ' \u2013 ' + vessel;
  html += '<div style="font-weight:bold;border-bottom:1px solid #ddd;padding-bottom:3px;margin-bottom:4px;">' + header + '</div>';
  // From
  html += '<div><b>From:</b> ' + fromLoc.name + '</div>';
  if (lastFrom) html += '<div style="padding-left:14px;">&nbsp;|-- ' + lastFrom + '</div>';
  // To
  html += '<div style="margin-top:3px;"><b>To:</b> ' + toLoc.name + '</div>';
  if (firstTo) html += '<div style="padding-left:14px;">&nbsp;|-- ' + firstTo + '</div>';
  html += '</div>';
  return html;
}

// \u2500\u2500 Draw route lines (one per container, never crossing containers) \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var routeLayers = {};   // key -> { lines: [polyline,...], stopSet: Set of loc indices }
var lineGroups  = {};   // "fi-ti" -> count (lane separation counter)
ROUTE_DATA.routes.forEach(function(route) {
  routeLayers[route.key] = { lines: [], stopSet: {} };
  for (var i = 0; i < route.stops.length - 1; i++) {
    var fi = route.stops[i], ti = route.stops[i + 1];
    var pairKey = Math.min(fi, ti) + '-' + Math.max(fi, ti);
    var offset = (lineGroups[pairKey] || 0) * 2;
    lineGroups[pairKey] = (lineGroups[pairKey] || 0) + 1;
    var fromLoc = ROUTE_DATA.locations[fi];
    var toLoc   = ROUTE_DATA.locations[ti];
    var vessel  = (route.vessels && i < route.vessels.length) ? route.vessels[i] : '';
    var tip     = buildContainerLineTooltip(route, fromLoc, toLoc, vessel);
    var line    = drawArrowLine([fromLoc.lat, fromLoc.lon], [toLoc.lat, toLoc.lon], route.color, offset, tip);
    if (line) routeLayers[route.key].lines.push(line);
  }
  route.stops.forEach(function(si) { routeLayers[route.key].stopSet[si] = true; });
});

// \u2500\u2500 Draw location dots (one per unique location) \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var locationMarkers = [];
ROUTE_DATA.locations.forEach(function(loc) {
  var marker = L.circleMarker([loc.lat, loc.lon], {
    radius: 7, color: '#ffffff', weight: 2,
    fillColor: '#2c3e50', fillOpacity: 0.9
  }).addTo(map);
  marker.bindPopup(buildContainerPopup(loc), {maxWidth: 440});
  marker.bindTooltip(loc.name, {permanent: true, className: 'stop-label', direction: 'top', offset: [0, -8]});
  locationMarkers.push(marker);
});

// \u2500\u2500 Fit bounds \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var allCoords = ROUTE_DATA.locations.map(function(l) { return [l.lat, l.lon]; });
if (allCoords.length > 0) {
  try { map.fitBounds(L.latLngBounds(allCoords), {padding: [60, 60], maxZoom: 10}); } catch(e) {}
}

// \u2500\u2500 Legend with hover highlight per container \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
var legendCtrl = L.control({position: 'bottomright'});
legendCtrl.onAdd = function() {
  var div = L.DomUtil.create('div', 'legend');
  var h = '<h4>Legend</h4>';
  ROUTE_DATA.routes.forEach(function(route) {
    h += '<div class="legend-item" data-key="' + route.key + '">'
      + '<div class="leg-swatch" style="background:' + route.color + '"></div>'
      + '<span>' + route.key + '</span></div>';
  });
  (ROUTE_DATA.highlight_regions || []).forEach(function(r) {
    var m = (r.color||'').match(/rgba?\\((\\d+),\\s*(\\d+),\\s*(\\d+)/);
    var hex = r.color || '#ffa500';
    if (m) hex = '#' + [+m[1],+m[2],+m[3]].map(function(v){return v.toString(16).padStart(2,'0');}).join('');
    h += '<div class="legend-item"><div class="leg-zone" style="background:' + hex
       + ';opacity:0.55;border:2px solid ' + hex + '"></div><span>' + r.name + '</span></div>';
  });
  h += '<div style="margin-top:8px;border-top:1px solid #ddd;padding-top:6px;">'
     + '<div class="legend-item" id="wz-toggle" title="Click to toggle war zone overlay" style="cursor:pointer;">'
     + '<div style="width:22px;height:11px;border:2px dashed #c0392b;background:rgba(231,76,60,0.18);flex-shrink:0;border-radius:2px;"></div>'
     + '<span style="color:#c0392b;font-weight:600;">&#9888; War Zones</span></div></div>';
  div.innerHTML = h;

  // Hover: highlight hovered container, dim all others
  div.querySelectorAll('[data-key]').forEach(function(item) {
    var hoverKey = item.getAttribute('data-key');
    item.addEventListener('mouseenter', function() {
      ROUTE_DATA.routes.forEach(function(r) {
        var rl = routeLayers[r.key];
        if (!rl) return;
        var dimmed = (r.key !== hoverKey);
        rl.lines.forEach(function(l) {
          l.setStyle({opacity: dimmed ? 0.10 : 1.0, weight: dimmed ? 1.0 : 3.0});
        });
      });
      // Dim location dots not visited by hovered container
      var hoverStops = routeLayers[hoverKey] ? routeLayers[hoverKey].stopSet : {};
      locationMarkers.forEach(function(m, mi) {
        var visited = !!hoverStops[mi];
        m.setStyle({fillOpacity: visited ? 1.0 : 0.10, opacity: visited ? 1.0 : 0.15});
      });
    });
    item.addEventListener('mouseleave', function() {
      ROUTE_DATA.routes.forEach(function(r) {
        var rl = routeLayers[r.key];
        if (!rl) return;
        rl.lines.forEach(function(l) { l.setStyle({opacity: 0.85, weight: 1.8}); });
      });
      locationMarkers.forEach(function(m) { m.setStyle({fillOpacity: 0.9, opacity: 1.0}); });
    });
  });

  // War zone toggle
  var wzToggle = div.querySelector('#wz-toggle');
  if (wzToggle) {
    wzToggle.addEventListener('click', function() {
      warZonesVisible = !warZonesVisible;
      warZoneLayers.forEach(function(l) {
        if (warZonesVisible) { map.addLayer(l); } else { map.removeLayer(l); }
      });
      wzToggle.style.opacity = warZonesVisible ? '1' : '0.4';
    });
  }
  L.DomEvent.disableScrollPropagation(div);
  return div;
};
legendCtrl.addTo(map);

} catch(e) {
  var d = document.getElementById('map-error');
  if (d) { d.style.display='block'; d.innerHTML='<b>Map Error:</b> ' + e.message; }
}
</script>
</body>
</html>
"""

    route_data_json = _json.dumps(data, ensure_ascii=False)
    escaped_title   = _html.escape(title)
    html_content = (
        _CONTAINER_ROUTE_HTML
        .replace("|||TITLE|||", escaped_title)
        .replace("|||ROUTE_DATA_JSON|||", route_data_json)
    )
    full_path = f"{file_store_path}/{file_id}_{title_slug}.html"
    with open(full_path, "w", encoding="utf-8") as fh:
        fh.write(html_content)
    return _file_meta(file_id, f"{title_slug}.html", "text/html", full_path)


def _plot_leaflet_map(
    title: str,
    data: dict[str, Any],
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Generate a Leaflet.js interactive map as a self-contained HTML file.

    Supports: markers, bubbles (sized circle markers), route polylines,
    directional arrow decorators, and filled polygon zone overlays.
    """
    import json as _json

    MAX_MAP_POINTS = 5000
    map_truncated = False

    ROUTE_COLORS = [
        "#E07B39", "#1F4788", "#2ECC71", "#E74C3C", "#9B59B6",
        "#F39C12", "#1ABC9C", "#E91E63", "#00BCD4", "#8BC34A",
    ]

    arrows = data.get("arrows", False)
    zones_raw = data.get("zones") or []

    # ── groups → routes conversion ────────────────────────────────────────
    # If caller passes flat lat/lon/labels + a "groups" array, auto-build
    # the routes list so they don't have to construct nested objects.
    # groups values may be composite "TrackNumber_ContainerNumber" keys —
    # each unique key becomes its own independent route (no cross-route connections).
    groups_raw       = data.get("groups")
    track_groups_raw = data.get("track_groups") or []
    # Single tracking number shorthand — when agent provides tracking="TRK" instead of
    # a full track_groups parallel array, fill track_groups_raw from it later.
    _single_tracking = str(data.get("tracking", "") or "").strip()

    import re as _grp_re
    _container_suffix_pat = _grp_re.compile(r'^[A-Z]{4}[0-9]{7}$')

    # ── Auto-derive groups from labels when agent omits the groups key ─────
    # Labels may be 'TrackNumber/ContainerNumber<br>...' or 'ContainerNumber<br>...'
    # Build composite 'TrackNumber_ContainerNumber' or plain container keys.
    if not groups_raw and data.get("labels"):
        _ag_trk_pat = _grp_re.compile(r'^(.+)/([A-Z]{4}[0-9]{7})')
        _ag_cont_pat = _grp_re.compile(r'^([A-Z]{4}[0-9]{7})')
        _auto_groups: list = []
        for _ag_lbl in data["labels"]:
            _ag_seg0 = (_ag_lbl.split("<br>")[0] if "<br>" in str(_ag_lbl) else str(_ag_lbl)).strip()
            _ag_m = _ag_trk_pat.match(_ag_seg0)
            if _ag_m:
                _auto_groups.append(f"{_ag_m.group(1)}_{_ag_m.group(2)}")
            else:
                _ag_m2 = _ag_cont_pat.match(_ag_seg0)
                _auto_groups.append(_ag_m2.group(1) if _ag_m2 else _ag_seg0)
        if _auto_groups:
            groups_raw = _auto_groups

    # Default arrows=True whenever groups are present (agent frequently omits this)
    if groups_raw and not data.get("arrows"):
        arrows = True

    def _grp_display_name(key: str) -> str:
        """For composite 'TRK_CONTAINER' keys, return just 'CONTAINER' for display."""
        if '_' in key:
            suffix = key.rsplit('_', 1)[-1]
            if _container_suffix_pat.match(suffix):
                return suffix
        return key

    if groups_raw and not data.get("routes"):
        raw_lats   = [float(v) for v in (data.get("lat")    or [])]
        raw_lons   = [float(v) for v in (data.get("lon")    or [])]
        raw_labels = list(data.get("labels") or [str(i) for i in range(len(raw_lats))])
        raw_sizes  = list(data.get("sizes")  or [])
        # Preserve insertion order; key is the full composite group string
        seen: dict[str, dict] = {}
        for idx, grp in enumerate(groups_raw):
            grp = str(grp)
            if grp not in seen:
                seen[grp] = {"lat": [], "lon": [], "labels": [], "sizes": []}
            if idx < len(raw_lats):
                seen[grp]["lat"].append(raw_lats[idx])
                seen[grp]["lon"].append(raw_lons[idx])
                seen[grp]["labels"].append(raw_labels[idx] if idx < len(raw_labels) else "")
                seen[grp]["sizes"].append(raw_sizes[idx] if idx < len(raw_sizes) else None)

        # ── Re-split mixed groups by container number ─────────────────────
        # If any group bucket holds stops for more than one container (e.g.
        # the agent used tracking numbers as group keys instead of composite
        # TrackNumber_ContainerNumber keys), split it automatically so each
        # container becomes its own independent route.
        import re as _split_re
        _lbl_cont_re = _split_re.compile(r'^[A-Z]{4}[0-9]{7}$')

        def _label_container(lbl: str) -> str:
            """Return the bare container number from a label's first <br> segment, or ''."""
            seg0 = lbl.split("<br>")[0] if "<br>" in lbl else lbl
            seg0 = _split_re.sub(r"<[^>]+>", "", seg0).strip()
            return seg0 if _lbl_cont_re.match(seg0) else ""

        _needs_split = any(
            len({_label_container(lbl) for lbl in gd["labels"]} - {""}) > 1
            for gd in seen.values()
        )
        if _needs_split:
            _new_seen: dict[str, dict] = {}
            for _gk, _gd in seen.items():
                for _i, _lbl in enumerate(_gd["labels"]):
                    _cont = _label_container(_lbl)
                    _sub_key = f"{_gk}_{_cont}" if _cont else _gk
                    if _sub_key not in _new_seen:
                        _new_seen[_sub_key] = {"lat": [], "lon": [], "labels": [], "sizes": []}
                    _new_seen[_sub_key]["lat"].append(_gd["lat"][_i])
                    _new_seen[_sub_key]["lon"].append(_gd["lon"][_i])
                    _new_seen[_sub_key]["labels"].append(_lbl)
                    _new_seen[_sub_key]["sizes"].append(
                        _gd["sizes"][_i] if _i < len(_gd["sizes"]) else None)
            seen = _new_seen

        # ── Color assignment: shade by tracking number ───────────────────
        # Composite keys (TrackNumber_ContainerNumber): one base hue per
        # tracking number, containers shaded light → dark within each hue.
        # Plain container keys: distinct color per container.
        import re as _cre
        _composite_pat = _cre.compile(r'^(.+)_([A-Z]{4}[0-9]{7})$')
        _all_grps = list(seen.keys())

        def _extract_track(key: str):
            m = _composite_pat.match(key)
            return m.group(1) if m else None

        # Base hues — one distinct saturated color per tracking number
        TRACK_BASE_COLORS = [
            "#27AE60",  # Green
            "#2980B9",  # Blue
            "#E67E22",  # Orange
            "#8E44AD",  # Purple
            "#C0392B",  # Red
            "#16A085",  # Teal
            "#D35400",  # Burnt Orange
            "#1A5276",  # Dark Navy
            "#6C3483",  # Violet
            "#1E8449",  # Dark Green
        ]

        def _blend_shade(hex_col: str, factor: float) -> str:
            """factor < 1.0 → mix with white (lighter); factor = 1.0 → base; factor > 1.0 → darken."""
            r = int(hex_col[1:3], 16)
            g = int(hex_col[3:5], 16)
            b = int(hex_col[5:7], 16)
            if factor <= 1.0:
                r = int(r * factor + 255 * (1 - factor))
                g = int(g * factor + 255 * (1 - factor))
                b = int(b * factor + 255 * (1 - factor))
            else:
                r = max(0, int(r * (2 - factor)))
                g = max(0, int(g * (2 - factor)))
                b = max(0, int(b * (2 - factor)))
            return "#%02x%02x%02x" % (min(255, r), min(255, g), min(255, b))

        _all_composite = bool(_all_grps) and all(_extract_track(g) for g in _all_grps)

        if _all_composite:
            # One base hue per unique tracking number (insertion order)
            _unique_tracks = list(dict.fromkeys(_extract_track(g) for g in _all_grps))
            _track_base = {t: TRACK_BASE_COLORS[i % len(TRACK_BASE_COLORS)]
                           for i, t in enumerate(_unique_tracks)}
            # Containers per tracking number (insertion order)
            _track_containers: dict[str, list] = {}
            for _g in _all_grps:
                _track_containers.setdefault(_extract_track(_g), []).append(_g)

            routes_raw = []
            for grp, v in seen.items():
                trk = _extract_track(grp)
                base = _track_base[trk]
                siblings = _track_containers[trk]
                n = len(siblings)
                idx_in_track = siblings.index(grp)
                # Light (factor=0.35) → dark (factor=1.1), evenly spread
                factor = 0.35 + (idx_in_track / max(n - 1, 1)) * 0.75 if n > 1 else 0.85
                routes_raw.append({
                    "name": _grp_display_name(grp), "tracking": trk,
                    "lat": v["lat"], "lon": v["lon"],
                    "labels": v["labels"], "sizes": [s for s in v["sizes"] if s is not None],
                    "color": _blend_shade(base, factor),
                })
        else:
            # Plain container keys — use track_groups_raw for tracking info and shading
            _cont_pat = _cre.compile(r'^[A-Z]{4}[0-9]{7}$')
            if _all_grps and all(_cont_pat.match(str(g)) for g in _all_grps):
                if track_groups_raw:
                    # Build container → tracking map (first occurrence wins)
                    _cont_to_trk: dict[str, str] = {}
                    for _ti, _gv in enumerate(groups_raw):
                        _gd = str(_gv)
                        if _ti < len(track_groups_raw) and _gd not in _cont_to_trk:
                            _cont_to_trk[_gd] = str(track_groups_raw[_ti])
                    _trk_unique2 = list(dict.fromkeys(_cont_to_trk.values()))
                    _trk_base2 = {t: TRACK_BASE_COLORS[i % len(TRACK_BASE_COLORS)]
                                  for i, t in enumerate(_trk_unique2)}
                    _trk_conts2: dict[str, list] = {}
                    for _ck in _all_grps:
                        _trk_conts2.setdefault(_cont_to_trk.get(_ck, ""), []).append(_ck)
                    routes_raw = []
                    for grp, v in seen.items():
                        _trk = _cont_to_trk.get(grp, "")
                        _base2 = _trk_base2.get(_trk, ROUTE_COLORS[0])
                        _sibs2 = _trk_conts2.get(_trk, [grp])
                        _n2 = len(_sibs2)
                        _idx2 = _sibs2.index(grp) if grp in _sibs2 else 0
                        _factor2 = 0.35 + (_idx2 / max(_n2 - 1, 1)) * 0.75 if _n2 > 1 else 0.85
                        routes_raw.append({
                            "name": grp, "tracking": _trk,
                            "lat": v["lat"], "lon": v["lon"],
                            "labels": v["labels"], "sizes": [s for s in v["sizes"] if s is not None],
                            "color": _blend_shade(_base2, _factor2),
                        })
                else:
                    # No track_groups from agent — use _single_tracking if provided
                    routes_raw = []
                    _st = _single_tracking  # may be "" if agent omitted tracking param too
                    _st_all = list(seen.keys())
                    _st_n = len(_st_all)
                    _st_base = TRACK_BASE_COLORS[0] if _st else None
                    for _idx, (grp, v) in enumerate(seen.items()):
                        if _st and _st_base:
                            _st_factor = 0.35 + (_idx / max(_st_n - 1, 1)) * 0.75 if _st_n > 1 else 0.85
                            _st_color = _blend_shade(_st_base, _st_factor)
                        else:
                            _st_color = ROUTE_COLORS[_idx % len(ROUTE_COLORS)]
                        routes_raw.append({
                            "name": grp, "tracking": _st,
                            "lat": v["lat"], "lon": v["lon"],
                            "labels": v["labels"], "sizes": [s for s in v["sizes"] if s is not None],
                            "color": _st_color,
                        })
            else:
                routes_raw = [
                    {"name": _grp_display_name(grp), "tracking": "",
                     "lat": v["lat"], "lon": v["lon"],
                     "labels": v["labels"], "sizes": [s for s in v["sizes"] if s is not None]}
                    for grp, v in seen.items()
                ]
    else:
        routes_raw = data.get("routes")

    map_data: dict[str, Any] = {
        "title": title,
        "routes": [],
        "zones": [],
        "highlight_regions": [],
        "arrows": arrows,
        "hide_route_tracking": bool(data.get("hide_route_tracking", False)),
    }

    all_lats: list[float] = []
    all_lons: list[float] = []

    # ── Zones ────────────────────────────────────────────────────────────
    for zone in zones_raw:
        z_lats = zone.get("lat", [])
        z_lons = zone.get("lon", [])
        if not z_lats or not z_lons:
            continue
        map_data["zones"].append({
            "name": zone.get("name", "Zone"),
            "latlngs": [[float(la), float(lo)] for la, lo in zip(z_lats, z_lons)],
            "color": zone.get("color", "rgba(255,0,0,0.25)"),
        })

    # ── Country / region highlights ──────────────────────────────────────
    for region in data.get("highlight_regions", []):
        map_data["highlight_regions"].append({
            "name": region.get("name", ""),
            "color": region.get("color", "rgba(31,71,136,0.25)"),
        })

    # ── Shared dedup helper ───────────────────────────────────────────────
    def _dedup_route(lats, lons, labels, sizes):
        """Collapse consecutive identical coordinates, keeping the longest label."""
        d_lats: list[float] = []
        d_lons: list[float] = []
        d_labels: list[str] = []
        d_sizes: list = []
        for idx, (la, lo) in enumerate(zip(lats, lons)):
            lbl = labels[idx] if idx < len(labels) else str(idx)
            sz  = sizes[idx]  if idx < len(sizes)  else None
            if d_lats and abs(la - d_lats[-1]) < 1e-6 and abs(lo - d_lons[-1]) < 1e-6:
                # Same location as previous — append new label as extra line
                d_labels[-1] = d_labels[-1] + "<br>" + lbl
            else:
                d_lats.append(la)
                d_lons.append(lo)
                d_labels.append(lbl)
                d_sizes.append(sz)
        return d_lats, d_lons, d_labels, [s for s in d_sizes if s is not None]

    # ── Routes ───────────────────────────────────────────────────────────
    if routes_raw:
        points_used = 0
        for i, route in enumerate(routes_raw):
            if points_used >= MAX_MAP_POINTS:
                map_truncated = True
                break
            r_lats = [float(v) for v in (route.get("lat") or [])]
            r_lons = [float(v) for v in (route.get("lon") or [])]
            r_labels = route.get("labels") or [str(j) for j in range(len(r_lats))]
            r_name = route.get("name", f"Route {i + 1}")
            r_connections = route.get("connections") or []
            r_sizes = route.get("sizes") or []

            remaining = MAX_MAP_POINTS - points_used
            if len(r_lats) > remaining:
                map_truncated = True
            r_lats = r_lats[:remaining]
            r_lons = r_lons[:remaining]
            r_labels = r_labels[:remaining]

            # Deduplicate consecutive identical coordinates within this route
            r_lats, r_lons, r_labels, r_sizes = _dedup_route(r_lats, r_lons, r_labels, r_sizes)

            # Rebuild connections for deduplicated points if not supplied
            if (arrows or r_connections) and len(r_lats) > 1:
                r_connections = [[j, j + 1] for j in range(len(r_lats) - 1)]

            color = route.get("color") or ROUTE_COLORS[i % len(ROUTE_COLORS)]
            map_data["routes"].append({
                "name": r_name,
                "tracking": route.get("tracking", ""),
                "color": color,
                "points": [[la, lo] for la, lo in zip(r_lats, r_lons)],
                "labels": r_labels,
                "connections": r_connections,
                "sizes": r_sizes,
            })
            all_lats.extend(r_lats)
            all_lons.extend(r_lons)
            points_used += len(r_lats)
    else:
        import re as _re
        raw_lats = [float(v) for v in (data.get("lat") or [])]
        raw_lons = [float(v) for v in (data.get("lon") or [])]
        if len(raw_lats) > MAX_MAP_POINTS:
            map_truncated = True
        lats = raw_lats[:MAX_MAP_POINTS]
        lons = raw_lons[:MAX_MAP_POINTS]
        labels = (data.get("labels") or [str(i) for i in range(len(lats))])[:MAX_MAP_POINTS]
        values = data.get("values") or []
        sizes  = data.get("sizes")  or []
        groups_flat = data.get("groups") or []

        # ── Auto-extract groups from label first line when groups not supplied ──
        # Labels are expected to start with the group name (container/track number)
        # followed by <br>.  E.g. "MSDU1234567<br>Houston<br>2026-02-09 (Actual)"
        # Extract the first <br>-delimited segment as the group key.
        if not groups_flat and labels:
            # Try to extract container/tracking number from first <br> segment.
            # A valid group key looks like a container number: 4 uppercase letters
            # followed by 7 digits (e.g. GAOU6335790) or a tracking number.
            _id_pat = _re.compile(r'^[A-Z0-9]{4,20}$')
            extracted = []
            for lbl in labels:
                norm = lbl.replace("\\n", "<br>").replace("\n", "<br>")
                first = _re.split(r"<br\s*/?>", norm, maxsplit=1)[0].strip()
                # Strip any HTML tags
                first = _re.sub(r"<[^>]+>", "", first).strip()
                extracted.append(first if first else "")
            # Use extracted groups only when all non-empty values look like IDs
            # (uppercase+digits) and there are multiple distinct values
            valid = [g for g in extracted if g]
            distinct = set(valid)
            if len(distinct) > 1 and all(_id_pat.match(g) for g in valid):
                groups_flat = extracted

        if groups_flat and len(groups_flat) == len(lats):
            # Build one route per unique group, preserving insertion order
            import collections as _col
            seen_grp: dict = _col.OrderedDict()
            for idx, grp in enumerate(groups_flat):
                grp = str(grp)
                if grp not in seen_grp:
                    seen_grp[grp] = {"lat": [], "lon": [], "labels": [], "sizes": []}
                if idx < len(lats):
                    seen_grp[grp]["lat"].append(lats[idx])
                    seen_grp[grp]["lon"].append(lons[idx])
                    seen_grp[grp]["labels"].append(labels[idx] if idx < len(labels) else "")
                    seen_grp[grp]["sizes"].append(sizes[idx] if idx < len(sizes) else None)

            # Shade containers by tracking number when all groups look like
            # container numbers (AAAA1234567) — same base colour, lighter → darker.
            # For multi-tracking without track_groups we cannot distinguish which
            # container belongs to which tracking, so shade all from ROUTE_COLORS[0].
            _p3_pat = _re.compile(r'^[A-Z]{4}[0-9]{7}$')
            _p3_grps = list(seen_grp.keys())
            _p3_all_containers = bool(_p3_grps) and all(
                _p3_pat.match(str(g)) for g in _p3_grps
            )
            _p3_n = len(_p3_grps)

            def _p3_shade(hex_col: str, factor: float) -> str:
                _r = int(hex_col[1:3], 16)
                _g = int(hex_col[3:5], 16)
                _b = int(hex_col[5:7], 16)
                return "#%02x%02x%02x" % (
                    max(0, min(255, int(_r * factor))),
                    max(0, min(255, int(_g * factor))),
                    max(0, min(255, int(_b * factor))),
                )

            for gi, (grp_name, gdata) in enumerate(seen_grp.items()):
                g_lats, g_lons, g_lbls, g_sizes = _dedup_route(
                    gdata["lat"], gdata["lon"], gdata["labels"],
                    [s for s in gdata["sizes"] if s is not None],
                )
                g_conn = [[j, j + 1] for j in range(len(g_lats) - 1)] if len(g_lats) > 1 else []
                color = ROUTE_COLORS[gi % len(ROUTE_COLORS)]
                map_data["routes"].append({
                    "name": grp_name,
                    "color": color,
                    "points": [[la, lo] for la, lo in zip(g_lats, g_lons)],
                    "labels": g_lbls,
                    "connections": g_conn,
                    "sizes": g_sizes,
                })
                all_lats.extend(g_lats)
                all_lons.extend(g_lons)
        else:
            # Single route — deduplicate consecutive identical coordinates
            dedup_lats: list[float] = []
            dedup_lons: list[float] = []
            dedup_labels: list[str] = []
            dedup_sizes: list = []
            for i, (la, lo) in enumerate(zip(lats, lons)):
                lbl = labels[i] if i < len(labels) else str(i)
                sz  = sizes[i]  if i < len(sizes)  else None
                if dedup_lats and abs(la - dedup_lats[-1]) < 1e-6 and abs(lo - dedup_lons[-1]) < 1e-6:
                    if len(lbl) > len(dedup_labels[-1]):
                        dedup_labels[-1] = lbl
                else:
                    dedup_lats.append(la); dedup_lons.append(lo)
                    dedup_labels.append(lbl); dedup_sizes.append(sz)

            lats, lons, labels = dedup_lats, dedup_lons, dedup_labels
            sizes = [s for s in dedup_sizes if s is not None]
            connections = data.get("connections") or []
            if arrows and not connections and len(lats) > 1:
                connections = [[i, i + 1] for i in range(len(lats) - 1)]
            if values and len(values) >= len(lats):
                labels = [f"{lb}<br>{v}" for lb, v in zip(labels, values)]
            map_data["routes"].append({
                "name": "",
                "color": ROUTE_COLORS[0],
                "points": [[la, lo] for la, lo in zip(lats, lons)],
                "labels": labels,
                "connections": connections,
                "sizes": sizes,
            })
            all_lats.extend(lats)
            all_lons.extend(lons)

    # ── Merge stop labels at shared locations across container routes ─────
    # When multiple containers share the same lat/lon stop, combine their
    # labels into one entry: LocationName<br>ContainerA  date  desc<br>...
    import re as _mlre
    _loc_stops: dict = {}   # (r_lat, r_lon) → [(route_idx, stop_idx, label)]
    for _ri, _route in enumerate(map_data["routes"]):
        for _si, _pt in enumerate(_route["points"]):
            _key = (round(_pt[0], 3), round(_pt[1], 3))
            _lbl = _route["labels"][_si] if _si < len(_route["labels"]) else ""
            _loc_stops.setdefault(_key, []).append((_ri, _si, _lbl))

    for _key, _stops in _loc_stops.items():
        if len(_stops) < 2:
            continue
        _loc_name = ""
        _container_parts: list = []   # list of (container_str, detail_str)
        for _ri, _si, _lbl in _stops:
            _segs = _mlre.split(r"<br\s*/?>", _lbl, flags=_mlre.IGNORECASE)
            _container = _segs[0].strip() if _segs else ""
            if not _loc_name and len(_segs) > 1:
                _loc_name = _segs[1].strip()
            # seg 2 = date, seg 3 = description
            _date = _segs[2].strip() if len(_segs) > 2 else ""
            _desc = _segs[3].strip() if len(_segs) > 3 else ""
            _detail = "  ".join(p for p in [_date, _desc] if p)
            _container_parts.append((_container, _detail))
        if _loc_name and _container_parts:
            # Location: bold, centred header
            _loc_html = f'<b style="display:block;text-align:center">{_loc_name}</b>'
            # Each container: numbered, name in bold
            _item_lines = []
            for _ci, (_cont, _det) in enumerate(_container_parts):
                _item = f'{_ci + 1}. <b>{_cont}</b>'
                if _det:
                    _item += f"  {_det}"
                _item_lines.append(_item)
            _combined = _loc_html + "<br>" + "<br>".join(_item_lines)
            for _ri, _si, _ in _stops:
                if _si < len(map_data["routes"][_ri]["labels"]):
                    map_data["routes"][_ri]["labels"][_si] = _combined

    # ── Reformat single-container labels: location first (bold/centred), name bold ──
    # Applies only to labels whose first segment is a standard container number
    # (4 uppercase letters + 7 digits).  Route-map tracking-number labels are
    # intentionally left unchanged.
    _sc_pat = _mlre.compile(r'^[A-Z]{4}[0-9]{7}$')
    for _ri, _route in enumerate(map_data["routes"]):
        for _si, _lbl in enumerate(_route["labels"]):
            _segs = _mlre.split(r"<br\s*/?>", _lbl, flags=_mlre.IGNORECASE)
            if len(_segs) >= 2 and _sc_pat.match(_segs[0].strip()):
                _cont = _segs[0].strip()
                _trk  = _route.get("tracking", "")
                _loc  = _segs[1].strip()
                _date = _segs[2].strip() if len(_segs) > 2 else ""
                _desc = _segs[3].strip() if len(_segs) > 3 else ""
                _detail = "  ".join(p for p in [_date, _desc] if p)
                _loc_html = f'<b style="display:block;text-align:center">{_loc}</b>'
                _item = f'<b data-trk="{_trk}">{_cont}</b>' if _trk else f'<b>{_cont}</b>'
                if _detail:
                    _item += f"  {_detail}"
                map_data["routes"][_ri]["labels"][_si] = _loc_html + "<br>" + _item

    # ── Lane offsets: separate parallel connections so they don't overlap ──
    # Two connections are "parallel" when their rounded endpoint coords match.
    # We spread them apart by ±LANE_PX pixels perpendicular to the route.
    _LANE_PX = 6
    _conn_reg: dict = {}
    for _ri, _route in enumerate(map_data["routes"]):
        _pts = _route["points"]
        for _ci, _conn in enumerate(_route.get("connections", [])):
            _fi, _ti = _conn[0], _conn[1]
            if _fi < len(_pts) and _ti < len(_pts):
                _fk = (round(_pts[_fi][0], 2), round(_pts[_fi][1], 2))
                _tk = (round(_pts[_ti][0], 2), round(_pts[_ti][1], 2))
                _conn_reg.setdefault((_fk, _tk), []).append((_ri, _ci))
    _offsets: dict = {_ri: [0] * len(_r.get("connections", []))
                      for _ri, _r in enumerate(map_data["routes"])}
    for _occs in _conn_reg.values():
        _n = len(_occs)
        if _n > 1:
            for _i, (_ri, _ci) in enumerate(_occs):
                _offsets[_ri][_ci] = (_i - (_n - 1) / 2.0) * _LANE_PX
    for _ri, _route in enumerate(map_data["routes"]):
        _route["lane_offsets"] = _offsets[_ri]

    # ── Centre / zoom ────────────────────────────────────────────────────
    if all_lats:
        center_lat = sum(all_lats) / len(all_lats)
        center_lon = sum(all_lons) / len(all_lons)
        spread = max(
            max(all_lats) - min(all_lats),
            max(all_lons) - min(all_lons),
        )
        zoom = 12 if spread < 0.5 else 9 if spread < 2 else 6 if spread < 10 else 4 if spread < 40 else 3 if spread < 100 else 2
    else:
        center_lat, center_lon, zoom = 20.0, 0.0, 2

    data_json = _json.dumps(map_data)

    # ── HTML template (uses ||| placeholders to avoid f-string conflicts) ─
    # Embed Leaflet CSS + JS inline so the map works without any CDN/internet.
    _vendor_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "vendor")
    try:
        with open(os.path.join(_vendor_dir, "leaflet.css"), encoding="utf-8") as _f:
            _leaflet_css = _f.read()
        with open(os.path.join(_vendor_dir, "leaflet.js"), encoding="utf-8") as _f:
            _leaflet_js = _f.read()
        with open(os.path.join(_vendor_dir, "topojson-client.min.js"), encoding="utf-8") as _f:
            _topojson_js = _f.read()
        _inline_head = (
            f"<style>\n{_leaflet_css}\n</style>\n"
            f"<script>\n{_leaflet_js}\n</script>\n"
            f"<script>\n{_topojson_js}\n</script>"
        )
        _cdn_tags = (
            '<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>\n'
            '<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>\n'
            '<script src="https://cdn.jsdelivr.net/npm/topojson-client@3/dist/topojson-client.min.js"></script>'
        )
    except Exception:
        _inline_head = None

    html = _LEAFLET_MAP_HTML
    if _inline_head:
        html = html.replace(_cdn_tags, _inline_head)
    html = html.replace("|||TITLE|||", title)
    html = html.replace("|||DATA_JSON|||", data_json)
    html = html.replace("|||CENTER_LAT|||", str(round(center_lat, 6)))
    html = html.replace("|||CENTER_LON|||", str(round(center_lon, 6)))
    html = html.replace("|||ZOOM|||", str(zoom))

    filename = f"{file_id}_{title_slug}.html"
    full_path = os.path.join(file_store_path, filename)
    with open(full_path, "w", encoding="utf-8") as fh:
        fh.write(html)

    meta = _file_meta(file_id, f"{title_slug}.html", "text/html", full_path)
    meta["map_truncated"] = map_truncated
    return meta


# -- interactive (plotly) ---------------------------------------------------


def _plot_interactive(
    plot_type: str,
    title: str,
    data: dict[str, Any],
    x_label: str | None,
    y_label: str | None,
    file_id: str,
    title_slug: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Render a chart with Plotly and save as self-contained HTML."""

    try:
        import plotly.graph_objects as go  # type: ignore[import-untyped]
    except ImportError:
        return {"error": "plotly is not installed. Run: pip install plotly"}

    labels = data.get("labels", [])
    values = data.get("values", [])
    fig: go.Figure | None = None

    if plot_type == "bar":
        fig = go.Figure(
            data=[go.Bar(x=labels, y=values, marker_color=HEADER_COLOR_HEX)]
        )

    elif plot_type == "line":
        fig = go.Figure(
            data=[go.Scatter(x=labels, y=values, mode="lines+markers",
                             line=dict(color=HEADER_COLOR_HEX))]
        )

    elif plot_type == "scatter":
        x = data.get("x", [])
        y = data.get("y", [])
        fig = go.Figure(
            data=[go.Scatter(x=x, y=y, mode="markers",
                             marker=dict(color=HEADER_COLOR_HEX, opacity=0.7))]
        )

    elif plot_type == "pie":
        # Align lengths, then cap at top 20 slices for readability
        _n = min(len(labels), len(values))
        _lbls, _vals = list(labels[:_n]), list(values[:_n])
        PIE_MAX = 20
        if _n > PIE_MAX:
            _paired = sorted(zip(_vals, _lbls), reverse=True)
            _vals  = [v for v, _ in _paired[:PIE_MAX - 1]] + [sum(v for v, _ in _paired[PIE_MAX - 1:])]
            _lbls  = [l for _, l in _paired[:PIE_MAX - 1]] + ["Other"]
        fig = go.Figure(
            data=[go.Pie(labels=_lbls, values=_vals, hole=0)]
        )

    elif plot_type == "heatmap":
        labels_x = data.get("labels_x", [])
        labels_y = data.get("labels_y", [])
        fig = go.Figure(
            data=[go.Heatmap(z=values, x=labels_x, y=labels_y,
                             colorscale="Blues")]
        )

    elif plot_type == "histogram":
        fig = go.Figure(
            data=[go.Histogram(x=values, marker_color=HEADER_COLOR_HEX)]
        )

    elif plot_type == "bar_stacked":
        PALETTE = [
            "#1F4788", "#E07B39", "#2ECC71", "#E74C3C", "#9B59B6",
            "#F39C12", "#1ABC9C", "#E91E63", "#00BCD4", "#8BC34A",
        ]
        series_list = data.get("series", [])
        traces = []
        for idx, s in enumerate(series_list):
            traces.append(go.Bar(
                name=s.get("name", f"Series {idx + 1}"),
                x=labels,
                y=s.get("values", []),
                marker_color=PALETTE[idx % len(PALETTE)],
            ))
        fig = go.Figure(data=traces)
        fig.update_layout(barmode="stack")

    elif plot_type == "map":
        return _plot_leaflet_map(
            title, data, file_id, title_slug, file_store_path,
        )

    else:
        return {"error": f"Unsupported plot_type: {plot_type}"}

    fig.update_layout(
        title=dict(text=title, font=dict(size=16)),
        xaxis_title=x_label or "",
        yaxis_title=y_label or "",
        template="plotly_white",
        margin=dict(l=60, r=30, t=60, b=60),
    )

    filename = f"{file_id}_{title_slug}.html"
    full_path = os.path.join(file_store_path, filename)
    fig.write_html(full_path, include_plotlyjs=True)

    return _file_meta(file_id, f"{title_slug}.html", "text/html", full_path)


# ---------------------------------------------------------------------------
# 2. generate_pdf
# ---------------------------------------------------------------------------

_PDF_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<style>
  @page {{
    size: A4 landscape;
    margin: 20mm 15mm 25mm 15mm;
    @bottom-center {{
      content: "Page " counter(page) " of " counter(pages);
      font-size: 9px;
      color: #888;
    }}
  }}
  body {{
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 11px;
    color: #222;
    line-height: 1.4;
  }}
  .report-header {{
    margin-bottom: 18px;
  }}
  .report-header h1 {{
    font-size: 22px;
    color: {header_color};
    margin: 0 0 4px 0;
  }}
  .report-header .timestamp {{
    font-size: 10px;
    color: #888;
  }}
  .summary {{
    margin-bottom: 16px;
    font-size: 12px;
    color: #444;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    font-size: 10px;
  }}
  th {{
    background-color: {header_color};
    color: #fff;
    font-weight: bold;
    text-align: left;
    padding: 6px 8px;
    border: 1px solid {header_color};
  }}
  td {{
    padding: 5px 8px;
    border: 1px solid #ddd;
  }}
  tr:nth-child(even) td {{
    background-color: {alt_row};
  }}
  td.num {{
    text-align: right;
    font-variant-numeric: tabular-nums;
  }}
</style>
</head>
<body>
  <div class="report-header">
    <h1>{title}</h1>
    <div class="timestamp">Generated {timestamp}</div>
  </div>
  {summary_html}
  <table>
    <thead><tr>{header_cells}</tr></thead>
    <tbody>{body_rows}</tbody>
  </table>
</body>
</html>
"""


def generate_pdf(
    title: str,
    columns: list[str],
    rows: list[list[str]],
    summary: str | None = None,
    filename: str | None = None,
    file_store_path: str = "./tmp/files",
) -> dict[str, Any]:
    """Generate a PDF report and save it to *file_store_path*.

    Attempts to use WeasyPrint for high-quality HTML->PDF conversion.
    Falls back to ReportLab if WeasyPrint is unavailable.
    """
    ensure_file_store(file_store_path)
    file_id = _short_uuid()
    safe_name = _slugify(filename) if filename else _slugify(title)

    try:
        return _pdf_weasyprint(
            title, columns, rows, summary, file_id, safe_name, file_store_path,
        )
    except ImportError:
        logger.info("WeasyPrint not available, falling back to ReportLab")
        try:
            return _pdf_reportlab(
                title, columns, rows, summary, file_id, safe_name, file_store_path,
            )
        except ImportError:
            return {
                "error": (
                    "Neither weasyprint nor reportlab is installed. "
                    "Install one of them: pip install weasyprint  or  pip install reportlab"
                )
            }
    except Exception as exc:
        logger.exception("generate_pdf failed")
        return {"error": str(exc)}


def _pdf_weasyprint(
    title: str,
    columns: list[str],
    rows: list[list[str]],
    summary: str | None,
    file_id: str,
    safe_name: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Render PDF with WeasyPrint (raises ImportError if not installed)."""
    from weasyprint import HTML  # type: ignore[import-untyped]

    # Detect which columns are numeric (by checking first data row).
    numeric_cols: set[int] = set()
    if rows:
        for idx, val in enumerate(rows[0]):
            if _is_numeric(val):
                numeric_cols.add(idx)

    header_cells = "".join(f"<th>{_esc(c)}</th>" for c in columns)

    body_parts: list[str] = []
    for row in rows:
        cells = []
        for idx, val in enumerate(row):
            cls = ' class="num"' if idx in numeric_cols else ""
            cells.append(f"<td{cls}>{_esc(str(val))}</td>")
        body_parts.append(f"<tr>{''.join(cells)}</tr>")

    summary_html = f'<div class="summary">{_esc(summary)}</div>' if summary else ""
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    html_str = _PDF_HTML_TEMPLATE.format(
        header_color=HEADER_COLOR_HEX,
        alt_row=ALT_ROW_COLOR,
        title=_esc(title),
        timestamp=timestamp,
        summary_html=summary_html,
        header_cells=header_cells,
        body_rows="\n".join(body_parts),
    )

    out_filename = f"{file_id}_{safe_name}.pdf"
    full_path = os.path.join(file_store_path, out_filename)
    HTML(string=html_str).write_pdf(full_path)

    return _file_meta(file_id, f"{safe_name}.pdf", "application/pdf", full_path)


def _pdf_reportlab(
    title: str,
    columns: list[str],
    rows: list[list[str]],
    summary: str | None,
    file_id: str,
    safe_name: str,
    file_store_path: str,
) -> dict[str, Any]:
    """Fallback PDF generation using ReportLab (raises ImportError if absent)."""
    from reportlab.lib import colors as rl_colors  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        SimpleDocTemplate,
        Table,
        TableStyle,
        Paragraph,
        Spacer,
    )

    out_filename = f"{file_id}_{safe_name}.pdf"
    full_path = os.path.join(file_store_path, out_filename)

    doc = SimpleDocTemplate(
        full_path,
        pagesize=landscape(A4),
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=20 * mm,
        bottomMargin=25 * mm,
    )

    styles = getSampleStyleSheet()
    elements: list[Any] = []

    # Title
    elements.append(Paragraph(title, styles["Title"]))
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    elements.append(Paragraph(f"Generated {timestamp}", styles["Normal"]))
    elements.append(Spacer(1, 6 * mm))

    # Summary
    if summary:
        elements.append(Paragraph(summary, styles["Normal"]))
        elements.append(Spacer(1, 4 * mm))

    # Table
    table_data = [columns] + [[str(v) for v in row] for row in rows]
    tbl = Table(table_data, repeatRows=1)

    # Derive header RGB from hex
    hr, hg, hb = (
        int(HEADER_COLOR_HEX[1:3], 16) / 255.0,
        int(HEADER_COLOR_HEX[3:5], 16) / 255.0,
        int(HEADER_COLOR_HEX[5:7], 16) / 255.0,
    )
    header_bg = rl_colors.Color(hr, hg, hb)

    style_commands: list[Any] = [
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.whitesmoke),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.Color(0.8, 0.8, 0.8)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ]
    # Alternating row colors
    for i in range(1, len(table_data)):
        if i % 2 == 0:
            style_commands.append(
                ("BACKGROUND", (0, i), (-1, i), rl_colors.Color(0.96, 0.96, 0.96))
            )

    tbl.setStyle(TableStyle(style_commands))
    elements.append(tbl)
    doc.build(elements)

    return _file_meta(file_id, f"{safe_name}.pdf", "application/pdf", full_path)


def _esc(text: str) -> str:
    """Minimal HTML escaping for template interpolation."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ---------------------------------------------------------------------------
# 3. generate_excel
# ---------------------------------------------------------------------------


def generate_excel(
    title: str,
    columns: list[str],
    rows: list[list[Any]],
    filename: str | None = None,
    file_store_path: str = "./tmp/files",
) -> dict[str, Any]:
    """Generate a formatted .xlsx workbook and save it to *file_store_path*.

    Styling matches user preferences from MEMORY.md:
      - Header: #1F4788 background, white bold text, font size 11
      - Freeze pane at row 1
      - Auto-sized columns (max 50 chars)
    """
    ensure_file_store(file_store_path)
    file_id = _short_uuid()
    safe_name = _slugify(filename) if filename else _slugify(title)

    try:
        from openpyxl import Workbook  # type: ignore[import-untyped]
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side  # type: ignore[import-untyped]
        from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]
    except ImportError:
        return {"error": "openpyxl is not installed. Run: pip install openpyxl"}

    try:
        wb = Workbook()
        ws = wb.active

        # Sheet name — Excel limits to 31 characters.
        ws.title = title[:31]

        # ── Header styling ────────────────────────────────────────────
        header_fill = PatternFill(
            start_color="1F4788", end_color="1F4788", fill_type="solid"
        )
        header_font = Font(
            name="Calibri", size=11, bold=True, color="FFFFFF"
        )
        header_alignment = Alignment(
            horizontal="center", vertical="center", wrap_text=True
        )
        thin_border = Border(
            left=Side(style="thin", color="CCCCCC"),
            right=Side(style="thin", color="CCCCCC"),
            top=Side(style="thin", color="CCCCCC"),
            bottom=Side(style="thin", color="CCCCCC"),
        )

        # ── Write header row ──────────────────────────────────────────
        for col_idx, col_name in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
            cell.border = thin_border

        # ── Data styling ──────────────────────────────────────────────
        data_font = Font(name="Calibri", size=11)
        alt_fill = PatternFill(
            start_color="F5F5F5", end_color="F5F5F5", fill_type="solid"
        )
        data_alignment = Alignment(vertical="center", wrap_text=False)

        # ── Write data rows ───────────────────────────────────────────
        for row_idx, row_data in enumerate(rows, start=2):
            for col_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.font = data_font
                cell.alignment = data_alignment
                cell.border = thin_border
                # Alternating row colors (even data rows = row_idx 3, 5, ...)
                if row_idx % 2 == 0:
                    cell.fill = alt_fill

        # ── Freeze top row ────────────────────────────────────────────
        ws.freeze_panes = "A2"

        # ── Auto-size columns ─────────────────────────────────────────
        for col_idx in range(1, len(columns) + 1):
            max_length = 0
            col_letter = get_column_letter(col_idx)
            for row_cells in ws.iter_rows(
                min_col=col_idx, max_col=col_idx,
                min_row=1, max_row=ws.max_row,
            ):
                for cell in row_cells:
                    val = str(cell.value) if cell.value is not None else ""
                    max_length = max(max_length, len(val))
            # Add a small buffer; cap at MAX_COL_WIDTH
            adjusted = min(max_length + 3, MAX_COL_WIDTH)
            ws.column_dimensions[col_letter].width = adjusted

        # ── Save ──────────────────────────────────────────────────────
        out_filename = f"{file_id}_{safe_name}.xlsx"
        full_path = os.path.join(file_store_path, out_filename)
        wb.save(full_path)

        return _file_meta(file_id, f"{safe_name}.xlsx",
                          "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                          full_path)

    except Exception as exc:
        logger.exception("generate_excel failed")
        return {"error": str(exc)}


# ---------------------------------------------------------------------------
# 4. generate_word_label
# ---------------------------------------------------------------------------


def generate_word_label(
    locations: list[dict[str, Any]],
    filename: str | None = None,
    file_store_path: str = "./tmp/files",
) -> dict[str, Any]:
    """Generate a Word (.docx) shipping label document.

    Each location becomes a bold/underlined header.  Under it, containers are
    numbered sequentially; the first event appears on the same line as the
    container name and subsequent events are indented to align with the date
    field of the first event.
    """
    ensure_file_store(file_store_path)
    file_id = _short_uuid()
    safe_name = _slugify(filename) if filename else "shipping_labels"

    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt, Inches  # type: ignore[import-untyped]
        from docx.oxml.ns import qn  # type: ignore[import-untyped]
        from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    except ImportError:
        return {
            "error": (
                "python-docx is not installed. "
                "Run: pip install python-docx"
            )
        }

    try:
        doc = Document()

        # Remove default styles' extra space and set body font to Calibri 11pt.
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Remove extra spacing on all paragraphs (Word adds 10pt after by default).
        from docx.shared import Pt as _Pt
        style.paragraph_format.space_before = _Pt(0)
        style.paragraph_format.space_after = _Pt(0)

        def _add_location_header(name: str, country_code: str, locode: str) -> None:
            """Add a bold+underline location header paragraph."""
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(f"{name}/{country_code} ({locode})")
            run.bold = True
            run.underline = True
            run.font.size = Pt(11)

        def _event_text(date: str, actual: bool, description: str) -> str:
            actual_marker = "(A)" if actual else "(E)"
            return f"{date} {actual_marker}: {description}"

        for location in locations:
            loc_name = location.get("name", "")
            country_code = location.get("country_code", "")
            locode = location.get("locode", "")
            containers = location.get("containers", [])

            _add_location_header(loc_name, country_code, locode)

            for seq, container in enumerate(containers, start=1):
                container_number = container.get("container_number", "")
                events = container.get("events", [])
                if not events:
                    continue

                first_event = events[0]
                first_event_text = _event_text(
                    first_event["date"], first_event["actual"], first_event["description"]
                )

                # ── First-event line: "{seq}.  ContainerName: date (A): desc" ──
                first_para = doc.add_paragraph()
                first_para.paragraph_format.space_before = Pt(2)
                first_para.paragraph_format.space_after = Pt(0)

                # Sequence number (normal weight)
                seq_run = first_para.add_run(f"{seq}.  ")
                seq_run.font.size = Pt(11)

                # Container number (bold)
                name_run = first_para.add_run(container_number)
                name_run.bold = True
                name_run.font.size = Pt(11)

                # Colon + first event (normal weight)
                event_run = first_para.add_run(f": {first_event_text}")
                event_run.font.size = Pt(11)

                # ── Continuation lines for remaining events ──────────────────
                # Estimate the left indent by approximating character widths.
                # prefix = "{seq}.  {container_number}: " — we use ~5.5pt per
                # char for normal text and ~6pt for bold, rounded up.
                prefix_normal = f"{seq}.  "   # normal weight chars
                prefix_bold = container_number  # bold chars
                prefix_colon = ": "             # normal weight
                indent_pt = (
                    len(prefix_normal) * 5.5
                    + len(prefix_bold) * 6.0
                    + len(prefix_colon) * 5.5
                )
                indent = Pt(indent_pt)

                for event in events[1:]:
                    cont_para = doc.add_paragraph()
                    cont_para.paragraph_format.space_before = Pt(0)
                    cont_para.paragraph_format.space_after = Pt(0)
                    cont_para.paragraph_format.left_indent = indent

                    cont_run = cont_para.add_run(
                        _event_text(event["date"], event["actual"], event["description"])
                    )
                    cont_run.font.size = Pt(11)

        out_filename = f"{file_id}_{safe_name}.docx"
        full_path = os.path.join(file_store_path, out_filename)
        doc.save(full_path)

        return _file_meta(
            file_id,
            f"{safe_name}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            full_path,
        )

    except Exception as exc:
        logger.exception("generate_word_label failed")
        return {"error": str(exc)}


# ---------------------------------------------------------------------------
# 5. cleanup_expired_files
# ---------------------------------------------------------------------------


def cleanup_expired_files(
    file_store_path: str = "./tmp/files",
    ttl_hours: int = 24,
) -> int:
    """Delete files older than *ttl_hours* from *file_store_path*.

    Returns the number of deleted files.
    """
    if not os.path.isdir(file_store_path):
        return 0

    now = time.time()
    cutoff = now - (ttl_hours * 3600)
    deleted = 0

    for entry in os.scandir(file_store_path):
        if entry.is_file():
            try:
                if entry.stat().st_mtime < cutoff:
                    os.remove(entry.path)
                    deleted += 1
                    logger.debug("Deleted expired file: %s", entry.name)
            except OSError as exc:
                logger.warning("Could not delete %s: %s", entry.path, exc)

    if deleted:
        logger.info(
            "Cleaned up %d expired file(s) from %s", deleted, file_store_path,
        )
    return deleted


# ---------------------------------------------------------------------------
# Convenience: dispatch a tool call by name
# ---------------------------------------------------------------------------


def handle_file_tool(tool_name: str, tool_input: dict[str, Any],
                     file_store_path: str = "./tmp/files") -> dict[str, Any]:
    """Dispatch a tool call to the appropriate generator function.

    This helper is used by the agent loop so it doesn't need to maintain
    its own if/elif chain for file-generation tools.
    """
    if tool_name == "generate_plot":
        return generate_plot(
            plot_type=tool_input["plot_type"],
            title=tool_input["title"],
            data=tool_input["data"],
            interactive=tool_input.get("interactive", False),
            x_label=tool_input.get("x_label"),
            y_label=tool_input.get("y_label"),
            file_store_path=file_store_path,
        )
    elif tool_name == "generate_pdf":
        return generate_pdf(
            title=tool_input["title"],
            columns=tool_input["columns"],
            rows=tool_input["rows"],
            summary=tool_input.get("summary"),
            filename=tool_input.get("filename"),
            file_store_path=file_store_path,
        )
    elif tool_name == "generate_excel":
        return generate_excel(
            title=tool_input["title"],
            columns=tool_input["columns"],
            rows=tool_input["rows"],
            filename=tool_input.get("filename"),
            file_store_path=file_store_path,
        )
    elif tool_name == "generate_word_label":
        return generate_word_label(
            locations=tool_input["locations"],
            filename=tool_input.get("filename"),
            file_store_path=file_store_path,
        )
    else:
        return {"error": f"Unknown file tool: {tool_name}"}
